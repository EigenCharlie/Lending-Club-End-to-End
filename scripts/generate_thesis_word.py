#!/usr/bin/env python3
"""
Generate complete thesis Word document from scratch.
APA 7 strict format + updated metrics (Venn-Abers) + auto TOC/LOT/LOF + embedded images.

Run: uv run --with python-docx scripts/generate_thesis_word.py
"""

import os
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

# ─── Paths ────────────────────────────────────────────────────────────────────
BASE = Path("/home/eigenlinux/projects/lending-club-risk-project")
INPUT = BASE / "thesis_poster" / "Documento_Tecnico_Final_Carlos_Vergara corrección.docx"
OUTPUT = BASE / "thesis_poster" / "Documento_Tecnico_Final_Carlos_Vergara_FINAL.docx"
IMG = BASE / "reports" / "notebook_images"

FIGURE_IMAGES = {
    3: str(IMG / "01_eda_lending_club" / "default_rate_by_grade.png"),
    4: str(IMG / "01_eda_lending_club" / "correlation_matrix.png"),
    5: str(IMG / "03_pd_modeling" / "roc_curves.png"),
    6: str(IMG / "03_pd_modeling" / "calibration_curves.png"),
    7: str(IMG / "03_pd_modeling" / "feature_importance.png"),
    8: str(IMG / "04_conformal_prediction" / "coverage_by_grade.png"),
    9: str(IMG / "04_conformal_prediction" / "interval_width_distribution.png"),
    10: str(IMG / "04_conformal_prediction" / "coverage_width_tradeoff.png"),
    11: str(IMG / "05_time_series_forecasting" / "ifrs9_scenario_fan_chart.png"),
}

# Table colors (hex, no #)
BLUE_HDR = "002060"
BLUE_LIGHT = "D9E1F2"
WHITE = "FFFFFF"

# ─── Text replacements (old → new) ────────────────────────────────────────────
REPS = [
    # Calibration method (ordered most-specific first)
    ("calibración con Isotonic Regression", "calibración Venn-Abers"),
    ("calibrado con Isotonic Regression", "calibrado con Venn-Abers"),
    ("calibrado con Isotonic regression", "calibrado con Venn-Abers"),
    ("Isotonic Regression", "Venn-Abers"),
    ("Isotonic regression", "Venn-Abers"),
    ("isotonic regression", "Venn-Abers"),
    ("Regresión Isotónica", "calibración Venn-Abers"),
    ("regresión isotónica", "calibración Venn-Abers"),
    ("calibración isotónica", "calibración Venn-Abers"),
    # AUC
    ("0.7116", "0.7130"),
    ("0.7115", "0.7130"),
    # Gini
    ("0.4230", "0.4260"),
    # Brier (table cells show "0.155" exactly)
    ("0.2313", "0.2313"),  # LR stays same
    ("0.1551", "0.1545"),
    ("0.155,", "0.1545,"),
    ("0.155 ", "0.1545 "),
    ("0.155\t", "0.1545\t"),
    # ECE
    ("0.0057", "0.0059"),
    ("5.93×10", "5.93×10"),  # keep as-is
    # KS
    ("0.3110", "0.3149"),
    ("0.311,", "0.3149,"),
    ("0.311 ", "0.3149 "),
    ("0.311\t", "0.3149\t"),
    # Coverage
    ("91.76%", "92.52%"),
    ("95.97%", "95.93%"),
    ("87.82%", "89.01%"),
    ("+29.13 pp", "+30.32 pp"),
    ("29.13 pp", "30.32 pp"),
    # Width
    ("0.752", "0.7546"),
    # ECL
    ("$977M", "$1,001M"),
    ("$1,791M", "$1,799M"),
    # Stats
    ("588 pruebas", "638 pruebas"),
    ("28 páginas", "30 páginas"),
]


def rep(text: str) -> str:
    """Apply all text replacements."""
    for old, new in REPS:
        text = text.replace(old, new)
    return text


# ─── XML helpers ──────────────────────────────────────────────────────────────
def make_toc_field(instruction: str) -> OxmlElement:
    """Create a Word TOC field paragraph (auto-updates on F9)."""
    p = OxmlElement("w:p")

    def _r(*children):
        r = OxmlElement("w:r")
        for c in children:
            r.append(c)
        return r

    fc_b = OxmlElement("w:fldChar")
    fc_b.set(qn("w:fldCharType"), "begin")
    fc_b.set(qn("w:dirty"), "true")

    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = instruction

    fc_s = OxmlElement("w:fldChar")
    fc_s.set(qn("w:fldCharType"), "separate")

    ph = OxmlElement("w:t")
    ph.text = "[Presione Ctrl+A → F9 para actualizar]"

    fc_e = OxmlElement("w:fldChar")
    fc_e.set(qn("w:fldCharType"), "end")

    p.append(_r(fc_b))
    p.append(_r(instr))
    p.append(_r(fc_s))
    p.append(_r(ph))
    p.append(_r(fc_e))
    return p


def make_page_break() -> OxmlElement:
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # Remove existing shd if any
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def add_custom_style(doc: Document, style_id: str, style_name: str):
    """Add a custom paragraph style based on Normal."""
    styles_el = doc.styles.element
    for s in styles_el.findall(qn("w:style")):
        if s.get(qn("w:styleId")) == style_id:
            return
    s = OxmlElement("w:style")
    s.set(qn("w:type"), "paragraph")
    s.set(qn("w:styleId"), style_id)
    nm = OxmlElement("w:name")
    nm.set(qn("w:val"), style_name)
    bo = OxmlElement("w:basedOn")
    bo.set(qn("w:val"), "Normal")
    s.append(nm)
    s.append(bo)
    styles_el.append(s)


def apply_style_id(para, style_id: str):
    """Apply a style ID to a paragraph's pPr."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._p.insert(0, pPr)
    ps = pPr.find(qn("w:pStyle"))
    if ps is None:
        ps = OxmlElement("w:pStyle")
        pPr.insert(0, ps)
    ps.set(qn("w:val"), style_id)


def is_image_para(para) -> bool:
    """Return True if paragraph contains an embedded image."""
    for elem in para._p.iter():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag in ("drawing", "pict", "object"):
            return True
    return False


# ─── APA style setup ──────────────────────────────────────────────────────────
def setup_styles(doc: Document):
    """Configure APA 7 styles on the document."""
    # Margins: 1 inch all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Normal style: Times New Roman 12pt, double spaced, 0.5" first indent
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = Pt(24)  # ~double (2 × 12pt)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    # Heading 1: bold, no indent, keep numbering from text
    for h_name in ("Heading 1", "Ttulo1", "Heading1"):
        try:
            h1 = doc.styles[h_name]
            h1.font.name = "Times New Roman"
            h1.font.size = Pt(12)
            h1.font.bold = True
            h1.paragraph_format.line_spacing = Pt(24)
            h1.paragraph_format.first_line_indent = Inches(0)
            h1.paragraph_format.space_before = Pt(0)
            h1.paragraph_format.space_after = Pt(0)
        except KeyError:
            pass

    # Heading 2: bold, no indent
    for h_name in ("Heading 2", "Ttulo2", "Heading2"):
        try:
            h2 = doc.styles[h_name]
            h2.font.name = "Times New Roman"
            h2.font.size = Pt(12)
            h2.font.bold = True
            h2.paragraph_format.line_spacing = Pt(24)
            h2.paragraph_format.first_line_indent = Inches(0)
            h2.paragraph_format.space_before = Pt(0)
            h2.paragraph_format.space_after = Pt(0)
        except KeyError:
            pass

    # Custom label styles (for auto LOT/LOF)
    add_custom_style(doc, "TablaLabel", "Tabla Label")
    add_custom_style(doc, "FiguraLabel", "Figura Label")


# ─── Paragraph helpers ────────────────────────────────────────────────────────
def add_para_centered(doc: Document, text: str, bold=False, size_pt=12) -> "Paragraph":
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    return p


def add_label_para(doc: Document, label: str, is_figura: bool) -> "Paragraph":
    """Add APA 7 figure/table label: bold, left, no italic, custom style."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(label)
    run.bold = True
    run.italic = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    style_id = "FiguraLabel" if is_figura else "TablaLabel"
    apply_style_id(p, style_id)
    return p


def add_title_para(doc: Document, title: str) -> "Paragraph":
    """Add APA 7 figure/table title: italic, left, no bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(rep(title))
    run.bold = False
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_figure_image(doc: Document, fig_num: int, extracted_images: dict):
    """Add figure image centered at 5.5 inches width."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run()

    img_path = FIGURE_IMAGES.get(fig_num)
    if img_path and os.path.exists(img_path):
        run.add_picture(img_path, width=Inches(5.5))
        return

    # Fallback: use extracted image from original docx
    name1 = f"image{fig_num}.jpeg"
    name2 = f"image{fig_num}.png"
    img_data = extracted_images.get(name1) or extracted_images.get(name2)
    if img_data:
        import io

        run.add_picture(io.BytesIO(img_data), width=Inches(5.5))
        return

    # Last resort: placeholder text
    run.text = f"[Figura {fig_num} — insertar imagen]"
    run.italic = True


def copy_runs(src_para, dst_para):
    """Copy runs from source paragraph to destination, applying replacements."""
    full_text = rep(src_para.text)
    if not full_text.strip():
        return
    # Simple: one run with the full text and formatting from first non-empty run
    bold = any(r.bold for r in src_para.runs if r.text.strip())
    italic = any(r.italic for r in src_para.runs if r.text.strip())
    run = dst_para.add_run(full_text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


# ─── Table reproduction ───────────────────────────────────────────────────────
def reproduce_table(src_table: DocxTable, doc: Document):
    """Reproduce a table with APA-style colored header."""
    nrows = len(src_table.rows)
    ncols = len(src_table.columns)
    if nrows == 0 or ncols == 0:
        return

    new_table = doc.add_table(rows=nrows, cols=ncols)
    new_table.style = "Table Grid"

    # Set column widths (equal distribution within 6.5 inches)
    col_width = Inches(6.5 / ncols)
    for col in new_table.columns:
        for cell in col.cells:
            cell.width = col_width

    for r_idx, src_row in enumerate(src_table.rows):
        for c_idx, src_cell in enumerate(src_row.cells):
            dst_cell = new_table.cell(r_idx, c_idx)
            cell_text = rep(src_cell.text.strip())

            # Set text
            dst_cell.paragraphs[0].clear()
            run = dst_cell.paragraphs[0].add_run(cell_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

            # Header row: dark blue background, white bold text
            if r_idx == 0:
                set_cell_bg(dst_cell, BLUE_HDR)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                # Alternate: light blue / white
                bg = BLUE_LIGHT if r_idx % 2 == 1 else WHITE
                set_cell_bg(dst_cell, bg)
                # Preserve bold from source
                src_run_bold = any(
                    r.bold for p in src_cell.paragraphs for r in p.runs if r.text.strip()
                )
                run.bold = src_run_bold

            # Cell paragraph formatting
            dst_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            dst_cell.paragraphs[0].paragraph_format.first_line_indent = Inches(0)
            dst_cell.paragraphs[0].paragraph_format.line_spacing = Pt(12)
            dst_cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            dst_cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    # Add empty paragraph after table for spacing
    doc.add_paragraph()


# ─── Portada ──────────────────────────────────────────────────────────────────
def add_portada(doc: Document, src_doc: Document):
    """Reproduce portada from source document (paras 0–12)."""
    # We manually rebuild it from the known structure
    doc.add_paragraph()  # top spacing

    # Title (centered, bold)
    add_para_centered(
        doc,
        "Conformal Prediction para la Calibración y Cuantificación de Incertidumbre "
        "en Modelos de Riesgo Crediticio",
        bold=True,
    )

    doc.add_paragraph()
    doc.add_paragraph()

    add_para_centered(doc, "Carlos Alfredo Vergara Rojas", bold=True)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    add_para_centered(doc, "Docente:", bold=True)
    add_para_centered(doc, "Alejandra María Restrepo Franco", bold=True)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    add_para_centered(
        doc,
        "UNIVERSIDAD TECNOLOGICA DE PEREIRA\n"
        "PROGRAMA DE ESPECIALIZACIÓN EN ANALÍTICA Y CIENCIA DE DATOS APLICADA\n"
        "COLOMBIA\n"
        "10 de marzo de 2026",
        bold=True,
    )


# ─── TOC / LOT / LOF sections ─────────────────────────────────────────────────
def add_index_sections(doc: Document):
    """Add CONTENIDO, LISTA DE TABLAS, LISTA DE FIGURAS with auto Word fields."""

    def _heading(text):
        p = doc.add_heading(text, level=1)
        p.paragraph_format.first_line_indent = Inches(0)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True

    # Page break → CONTENIDO
    doc.element.body.append(make_page_break())
    _heading("CONTENIDO")
    doc.element.body.append(make_toc_field(' TOC \\o "1-2" \\h \\z \\u '))
    doc.add_paragraph()

    # Page break → LISTA DE TABLAS
    doc.element.body.append(make_page_break())
    _heading("LISTA DE TABLAS")
    doc.element.body.append(make_toc_field(' TOC \\h \\z \\t "Tabla Label,1" '))
    doc.add_paragraph()

    # Page break → LISTA DE FIGURAS
    doc.element.body.append(make_page_break())
    _heading("LISTA DE FIGURAS")
    doc.element.body.append(make_toc_field(' TOC \\h \\z \\t "Figura Label,1" '))
    doc.add_paragraph()


# ─── Label detection regexes ──────────────────────────────────────────────────
LABEL_COMB_RE = re.compile(r"^((?:Tabla|Figura)\s+\d+)\.?\s+(.+)", re.DOTALL | re.IGNORECASE)
LABEL_ONLY_RE = re.compile(r"^((?:Tabla|Figura)\s+\d+)\.?\s*$", re.IGNORECASE)


def parse_label(text: str):
    """Return (label_str, title_str, fig_num, is_figura) or None."""
    m = LABEL_COMB_RE.match(text)
    if m:
        label = m.group(1).strip()
        title = m.group(2).strip()
        num = int(re.search(r"\d+", label).group())
        return label, title, num, label.lower().startswith("figura")
    m = LABEL_ONLY_RE.match(text)
    if m:
        label = m.group(1).strip().rstrip(".")
        num = int(re.search(r"\d+", label).group())
        return label, None, num, label.lower().startswith("figura")
    return None


# ─── Main body processing ─────────────────────────────────────────────────────
def process_body(doc: Document, src_doc: Document, extracted_images: dict):
    """
    Iterate source document body (from ANEXO onwards) and reproduce in new doc.
    Handles: headings, normal text, figure/table labels + images, tables.
    """
    body_elems = list(src_doc.element.body)

    # Find start index: first element after CONTENIDO section (ANEXO heading)
    start_idx = None
    for i, elem in enumerate(body_elems):
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "p":
            para = Paragraph(elem, src_doc)
            if "ANEXO: DECLARACION" in para.text.upper():
                start_idx = i
                break

    if start_idx is None:
        print("WARNING: Could not find ANEXO heading, starting from para 17")
        start_idx = 17

    # Page break before ANEXO
    doc.element.body.append(make_page_break())

    # State tracking
    last_fig_num = None
    after_label = False  # label-only just processed, next para is title
    after_title = False  # title just processed, next image para is figure

    for i in range(start_idx, len(body_elems)):
        elem = body_elems[i]
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

        # ── TABLE ──────────────────────────────────────────────────────────────
        if tag == "tbl":
            tbl = DocxTable(elem, src_doc)
            reproduce_table(tbl, doc)
            after_label = False
            after_title = False
            last_fig_num = None
            continue

        if tag != "p":
            continue

        # ── PARAGRAPH ──────────────────────────────────────────────────────────
        para = Paragraph(elem, src_doc)
        text = para.text.strip()
        style = para.style.name

        # Skip empty portada elements that leaked (just in case)
        if i < start_idx:
            continue

        # ── Image paragraph ────────────────────────────────────────────────────
        if is_image_para(para):
            if after_title and last_fig_num is not None:
                add_figure_image(doc, last_fig_num, extracted_images)
            # else: skip orphaned image paragraphs
            after_title = False
            last_fig_num = None
            continue

        # ── Figure/Table label ────────────────────────────────────────────────
        parsed = parse_label(text)
        if parsed:
            label, title, num, is_fig = parsed

            if title:
                # Combined label+title in one paragraph
                add_label_para(doc, label, is_fig)
                add_title_para(doc, title)
                if is_fig:
                    last_fig_num = num
                    after_title = True
                    after_label = False
                else:
                    last_fig_num = None
                    after_title = False
                    after_label = False
            else:
                # Label only → next para is the title
                add_label_para(doc, label, is_fig)
                if is_fig:
                    last_fig_num = num
                    after_label = True
                    after_title = False
                else:
                    last_fig_num = None
                    after_label = False
                    after_title = False
            continue

        # ── Title paragraph (right after label-only) ──────────────────────────
        if after_label:
            if text:
                add_title_para(doc, text)
            after_label = False
            after_title = last_fig_num is not None
            continue

        # Reset figure tracking if we hit a non-image, non-empty para after title
        if after_title and text:
            after_title = False
            last_fig_num = None

        # ── Headings ──────────────────────────────────────────────────────────
        if any(s in style for s in ("Heading 1", "Ttulo1", "Heading1")):
            p = doc.add_heading(rep(text), level=1)
            p.paragraph_format.first_line_indent = Inches(0)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = True
            continue

        if any(s in style for s in ("Heading 2", "Ttulo2", "Heading2")):
            p = doc.add_heading(rep(text), level=2)
            p.paragraph_format.first_line_indent = Inches(0)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = True
            continue

        # ── Empty paragraph ───────────────────────────────────────────────────
        if not text:
            ep = doc.add_paragraph()
            ep.paragraph_format.line_spacing = Pt(12)
            ep.paragraph_format.first_line_indent = Inches(0)
            continue

        # ── Normal body paragraph ─────────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = Pt(24)

        # Copy runs preserving bold/italic
        for src_run in para.runs:
            run_text = rep(src_run.text)
            if not run_text:
                continue
            run = p.add_run(run_text)
            run.bold = src_run.bold
            run.italic = src_run.italic
            run.underline = src_run.underline
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        # If no runs were added (all empty), add the full text
        if not p.runs and text:
            run = p.add_run(rep(text))
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


# ─── Main ─────────────────────────────────────────────────────────────────────
def main():
    print(f"Input : {INPUT}")
    print(f"Output: {OUTPUT}")

    # Load source
    src_doc = Document(str(INPUT))

    # Extract embedded images from source docx (for Figura 1, 2)
    extracted_images = {}
    with zipfile.ZipFile(str(INPUT)) as z:
        media = [f for f in z.namelist() if "word/media/" in f]
        for m in media:
            fname = os.path.basename(m)
            extracted_images[fname] = z.read(m)
    print(f"Extracted {len(extracted_images)} embedded images from source")

    # Map image1/image2 to figura 1/2
    # Sort media files to get consistent ordering
    media_sorted = sorted(extracted_images.keys())
    if len(media_sorted) >= 1:
        FIGURE_IMAGES[1] = f"__embedded__{media_sorted[0]}"
    if len(media_sorted) >= 2:
        FIGURE_IMAGES[2] = f"__embedded__{media_sorted[1]}"

    # Patch add_figure_image to use extracted images for 1,2
    def _add_figure_image_patched(doc, fig_num, extracted):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run()

        img_key = FIGURE_IMAGES.get(fig_num, "")
        if isinstance(img_key, str) and img_key.startswith("__embedded__"):
            name = img_key[len("__embedded__") :]
            data = extracted.get(name)
            if data:
                import io

                try:
                    run.add_picture(io.BytesIO(data), width=Inches(5.5))
                    return
                except Exception as e:
                    print(f"  Warning: could not embed {name}: {e}")
        elif img_key and os.path.exists(img_key):
            try:
                run.add_picture(img_key, width=Inches(5.5))
                return
            except Exception as e:
                print(f"  Warning: could not embed {img_key}: {e}")

        run.text = f"[Figura {fig_num} — imagen no disponible]"
        run.italic = True

    # Monkey-patch for this run
    import sys

    _mod = sys.modules[__name__]
    _mod.add_figure_image = _add_figure_image_patched

    # Create new document
    new_doc = Document()
    setup_styles(new_doc)
    add_custom_style(new_doc, "TablaLabel", "Tabla Label")
    add_custom_style(new_doc, "FiguraLabel", "Figura Label")

    # 1. Portada
    print("Building portada...")
    add_portada(new_doc, src_doc)

    # 2. CONTENIDO + LISTA DE TABLAS + LISTA DE FIGURAS
    print("Adding index sections (TOC fields)...")
    add_index_sections(new_doc)

    # 3. Body (ANEXO onwards)
    print("Processing document body...")
    process_body(new_doc, src_doc, extracted_images)

    # Save
    new_doc.save(str(OUTPUT))
    print(f"\n✅ Saved → {OUTPUT}")

    # ─── Verification ────────────────────────────────────────────────────────
    print("\n=== VERIFICATION ===")
    check = Document(str(OUTPUT))
    headings = [
        p.text for p in check.paragraphs if "Heading" in p.style.name or "Ttulo" in p.style.name
    ]
    tabla_labels = [
        p
        for p in check.paragraphs
        if p._p.find(qn("w:pPr")) is not None
        and p._p.find(qn("w:pPr")).find(qn("w:pStyle")) is not None
        and p._p.find(qn("w:pPr")).find(qn("w:pStyle")).get(qn("w:val")) == "TablaLabel"
    ]
    figura_labels = [
        p
        for p in check.paragraphs
        if p._p.find(qn("w:pPr")) is not None
        and p._p.find(qn("w:pPr")).find(qn("w:pStyle")) is not None
        and p._p.find(qn("w:pPr")).find(qn("w:pStyle")).get(qn("w:val")) == "FiguraLabel"
    ]
    venn_count = sum(1 for p in check.paragraphs if "Venn-Abers" in p.text)
    toc_fields = sum(1 for p in check.paragraphs if "[Presione Ctrl+A" in p.text)

    print(f"  Headings found   : {len(headings)}")
    print(f"  TablaLabel paras : {len(tabla_labels)}  (expected 12)")
    print(f"  FiguraLabel paras: {len(figura_labels)} (expected 11)")
    print(f"  'Venn-Abers' hits: {venn_count}  (expected ≥3)")
    print(f"  TOC field paras  : {toc_fields} (expected 3)")
    print(f"  Total paragraphs : {len(check.paragraphs)}")
    print(f"  Total tables     : {len(check.tables)}")
    print("\nOpen in Word → Ctrl+A → F9 to populate all indices with page numbers.")


if __name__ == "__main__":
    main()
