#!/usr/bin/env python3
"""Generate the Documento Tecnico Final for the specialization thesis.

Scope: Conformal Prediction for calibration and uncertainty quantification
in credit risk models (PD, LGD, EAD). Mondrian conformal, comparative
evaluation, and IFRS9 regulatory impact through CP.

OUT OF SCOPE: portfolio optimization, causal inference, MLOps, fairness deep-dive.

Usage:
    uv run python scripts/generate_thesis_document.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── Paths ────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
FIGURES = ROOT / "thesis_poster" / "figures"
NB_IMAGES = ROOT / "reports" / "notebook_images"
OUTPUT_DIR = ROOT / "thesis_poster"
DOCX_PATH = OUTPUT_DIR / "Documento_Tecnico_Final_Carlos_Vergara.docx"

# ── Colors ───────────────────────────────────────────────────────────
UTP_BLUE = RGBColor(0x0B, 0x5E, 0xD7)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
TABLE_HEADER_BG = "0B5ED7"
TABLE_ALT_BG = "F0F4FA"


# ── Helpers ──────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "999999")
        borders.append(element)
    tblPr.append(borders)


def styled_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = DARK_GRAY


def add_paragraph(
    doc: Document,
    text: str,
    bold: bool = False,
    italic: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after: int = 6,
    first_line_indent: float = 0.0,
) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent > 0:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic


def add_figure(
    doc: Document, image_path: Path, caption: str, fig_num: int, width_inches: float = 5.5
) -> int:
    if not image_path.exists():
        add_paragraph(
            doc,
            f"[Figura {fig_num}: {caption} — archivo no encontrado: {image_path.name}]",
            italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        return fig_num + 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f"Figura {fig_num}. {caption}")
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    r.italic = True
    r.font.color.rgb = DARK_GRAY
    return fig_num + 1


def add_styled_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    caption: str = "",
    table_num: int = 0,
    footnote: str = "",
) -> None:
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        r = cap.add_run(f"Tabla {table_num}. {caption}")
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"
        r.italic = True
        r.font.color.rgb = DARK_GRAY

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, TABLE_HEADER_BG)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
            if i % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    if footnote:
        fn = doc.add_paragraph()
        fn.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fn.paragraph_format.space_after = Pt(4)
        r = fn.add_run(footnote)
        r.font.size = Pt(9)
        r.font.name = "Times New Roman"
        r.font.color.rgb = DARK_GRAY
        r.italic = True
    else:
        doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════
#  DOCUMENT SECTIONS
# ══════════════════════════════════════════════════════════════════════


def build_cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    logo_path = FIGURES / "logo-utp.jpg"
    if logo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo_path), width=Inches(1.5))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Conformal Prediction para la Calibracion y Cuantificacion "
        "de Incertidumbre en Modelos de Riesgo Crediticio"
    )
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.name = "Times New Roman"
    r.font.color.rgb = DARK_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Carlos Alfredo Vergara Rojas")
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"
    r.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Documento de investigacion para optar al titulo de Especialista\n"
        "en Analitica y Ciencia de Datos Aplicada"
    )
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    r.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Docente:\nAlejandra Maria Restrepo Franco")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    r.font.color.rgb = DARK_GRAY

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "UNIVERSIDAD TECNOLOGICA DE PEREIRA\n"
        "PROGRAMA DE ESPECIALIZACION EN ANALITICA Y\n"
        "CIENCIA DE DATOS, COLOMBIA\n"
        "Marzo de 2026"
    )
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.name = "Times New Roman"
    r.font.color.rgb = DARK_GRAY

    doc.add_page_break()


def build_abstract(doc: Document) -> None:
    styled_heading(doc, "RESUMEN", level=1)

    add_paragraph(
        doc,
        (
            "Los modelos de riesgo crediticio utilizados en el sector financiero tipicamente "
            "producen predicciones puntuales de probabilidad de incumplimiento (PD), perdida dado "
            "incumplimiento (LGD) y exposicion en caso de incumplimiento (EAD), sin cuantificar "
            "la incertidumbre asociada a dichas estimaciones. Esta limitacion compromete la "
            "robustez de las provisiones bajo IFRS 9 y dificulta la auditabilidad de los modelos "
            "ante reguladores. La presente investigacion implementa y evalua tecnicas de prediccion "
            "conformal (Conformal Prediction, CP) para abordar esta brecha, utilizando el dataset "
            "publico de LendingClub (2007-2020) con separacion temporal estricta out-of-time."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Se construyo un modelo base de PD mediante CatBoost con restricciones monotonicas y "
            "calibracion post-hoc Venn-Abers, alcanzando un AUC-ROC de 0.7127 y un ECE de 0.0067. "
            "Sobre este modelo calibrado se aplico prediccion conformal Mondrian por grado de riesgo "
            "(A-G) utilizando MAPIE 1.3.0, obteniendo una cobertura empirica del 92.42% al nivel "
            "nominal del 90%, con una cobertura minima por grado de 89.01%, lo que representa una "
            "mejora de +30 puntos porcentuales frente al enfoque Split Conformal global (58.69%). "
            "La extension a LGD (90.50% de cobertura con variante adaptativa) y EAD (91.20%) "
            "demostro la aplicabilidad de CP a la triada completa del riesgo crediticio, cerrando "
            "un vacio en la literatura existente."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "El impacto regulatorio se cuantifico mediante simulacion de perdidas crediticias "
            "esperadas (ECL) bajo cuatro escenarios, revelando un rango de $1,001M a $1,799M "
            "(+83.3%), que explicita la incertidumbre en las provisiones. Se propuso ademas el "
            "ancho del intervalo conformal como senal complementaria de SICR para la clasificacion "
            "por etapas IFRS 9. Los resultados confirman que la prediccion conformal es una "
            "herramienta viable, complementaria a la calibracion tradicional, y con potencial "
            "de adopcion en entornos bancarios reales."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        "Palabras clave: prediccion conformal, riesgo crediticio, calibracion de probabilidades, "
        "IFRS 9, cuantificacion de incertidumbre, machine learning.",
        bold=True,
        space_after=12,
    )

    doc.add_paragraph()
    styled_heading(doc, "ABSTRACT", level=1)

    add_paragraph(
        doc,
        (
            "Credit risk models typically produce point estimates of Probability of Default (PD), "
            "Loss Given Default (LGD), and Exposure at Default (EAD), without quantifying the "
            "associated uncertainty. This limitation undermines the robustness of IFRS 9 provisions "
            "and hinders model auditability. This research implements and evaluates Conformal "
            "Prediction (CP) techniques to address this gap, using the LendingClub public dataset "
            "(2007-2020) with strict out-of-time temporal splits."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "A PD base model was built using CatBoost with monotonic constraints and Venn-Abers "
            "post-hoc calibration (AUC-ROC: 0.7127, ECE: 0.0067). Mondrian conformal prediction "
            "by risk grade (A-G) achieved 92.42% empirical coverage at the 90% nominal level, with "
            "a minimum per-grade coverage of 89.01% — a +30 percentage point improvement over global "
            "Split Conformal (58.69%). Extension to LGD (90.50% coverage) and EAD (91.20%) "
            "demonstrated applicability to the full credit risk triad. Regulatory impact was "
            "quantified through ECL scenario simulation, revealing a range of $1,001M to $1,799M "
            "(+83.3%), and the conformal interval width was proposed as a complementary SICR signal "
            "for IFRS 9 stage classification."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        "Keywords: conformal prediction, credit risk, probability calibration, IFRS 9, "
        "uncertainty quantification, machine learning.",
        bold=True,
        space_after=12,
    )

    doc.add_page_break()


def build_table_of_contents(doc: Document) -> None:
    styled_heading(doc, "CONTENIDO", level=1)
    items = [
        "Resumen / Abstract",
        "1. Introduccion",
        "2. Planteamiento del Problema",
        "3. Justificacion de la Investigacion",
        "4. Objetivos",
        "5. Marco Teorico y Estado del Arte",
        "6. Metodologia",
        "   6.7 Respuesta a la Evaluacion del Anteproyecto",
        "   6.8 Cambios Metodologicos respecto al Anteproyecto",
        "7. Consideraciones Eticas",
        "8. Resultados",
        "   8.1 Analisis Exploratorio de Datos",
        "   8.2 Desempeno del Modelo de PD",
        "   8.3 Intervalos Conformales para PD (Mondrian)",
        "   8.4 Comparacion: Mondrian vs. Split Global",
        "   8.5 Intervalos Conformales para LGD y EAD",
        "   8.6 Impacto Regulatorio: IFRS 9 y ECL",
        "9. Discusion de Resultados",
        "   9.1 Hallazgos Principales",
        "   9.2 Resultados en el Marco CRISP-DM",
        "   9.3 Cumplimiento de Objetivos",
        "10. Conclusiones",
        "11. Lineas Futuras del Proyecto",
        "12. Referencias",
        "Anexo: Declaracion de uso de herramientas de inteligencia artificial",
    ]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("\n".join(items))
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    r.font.color.rgb = DARK_GRAY
    doc.add_page_break()


# ── 1. INTRODUCCION ──────────────────────────────────────────────────


def build_introduction(doc: Document) -> None:
    styled_heading(doc, "1. INTRODUCCION", level=1)

    add_paragraph(
        doc,
        (
            "En la actualidad, los sistemas financieros enfrentan un entorno complejo y altamente "
            "regulado, en el cual la correcta estimacion del riesgo crediticio constituye un pilar "
            "fundamental para la estabilidad economica y la confianza de los inversionistas. La "
            "probabilidad de incumplimiento (PD, por sus siglas en ingles), la perdida dado "
            "incumplimiento (LGD) y la exposicion en caso de incumplimiento (EAD) son metricas "
            "centrales utilizadas para calcular provisiones, asignar capital y cumplir con estandares "
            "regulatorios como la Norma Internacional de Informacion Financiera 9 (IFRS 9) y el marco "
            "regulatorio de Basilea III [4], [9]. Estas tres metricas alimentan directamente el calculo "
            "de la perdida esperada (Expected Credit Loss, ECL = PD x LGD x EAD), que determina la "
            "salud financiera de las entidades bancarias y su capacidad para absorber perdidas "
            "inesperadas."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Sin embargo, los modelos tradicionales de riesgo crediticio, tanto los clasicos "
            "(regresion logistica, modelos lineales generalizados) como los modernos basados en "
            "machine learning, presentan una limitacion critica: tipicamente ofrecen predicciones "
            "puntuales sin acompanarlas de una medida fiable de la incertidumbre asociada. Una "
            "prediccion de PD del 12% no indica si el modelo esta seguro de esa estimacion o si "
            "el valor real podria estar entre 8% y 17%. Esta ausencia de cuantificacion de "
            "incertidumbre tiene consecuencias directas: perdidas inesperadas por subestimacion del "
            "riesgo, provisiones excesivas por sobreestimacion conservadora, y desconfianza de "
            "auditores y reguladores que exigen mayor transparencia y robustez en los modelos "
            "[12], [15]."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "La prediccion conformal (Conformal Prediction, CP) surge como una alternativa "
            "metodologica rigurosa para abordar esta limitacion. Propuesta originalmente por Vovk, "
            "Gammerman y Shafer [18], CP es un marco estadistico que permite acompanar cada "
            "prediccion con intervalos de confianza que poseen garantias formales de cobertura, sin "
            "requerir supuestos parametricos sobre la distribucion de los datos. A diferencia de los "
            "metodos bayesianos (que dependen de distribuciones a priori) o del bootstrap (que carece "
            "de garantias formales en muestras finitas), CP ofrece una cobertura marginal garantizada "
            "bajo el unico supuesto de intercambiabilidad de los datos [1]. "
            "Esta propiedad hace de CP una tecnica especialmente prometedora para el sector financiero, "
            "donde la robustez estadistica y la transparencia metodologica son requisitos regulatorios "
            "fundamentales."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "La presente investigacion implementa y evalua tecnicas de Conformal Prediction en "
            "modelos de riesgo crediticio, con el proposito de mejorar la calibracion de las "
            "probabilidades predichas, cuantificar de manera explicita la incertidumbre asociada "
            "a PD, LGD y EAD, y demostrar su utilidad practica en el contexto regulatorio de IFRS 9. "
            "El estudio se apoya en el dataset publico de LendingClub (2007-2020) [11], uno de los "
            "conjuntos de datos mas completos en la literatura de riesgo crediticio, permitiendo "
            "replicabilidad y comparacion con trabajos previos. Se utiliza la "
            "variante Mondrian de prediccion conformal, que calcula intervalos por subgrupo (grado "
            "de riesgo), garantizando cobertura condicional por segmento y no solo cobertura promedio "
            "global. Los resultados demuestran coberturas empiricas que superan los niveles nominales "
            "objetivo, validando la efectividad del enfoque propuesto."
        ),
        first_line_indent=1.25,
    )


# ── 2. PLANTEAMIENTO DEL PROBLEMA ────────────────────────────────────


def build_problem_statement(doc: Document) -> None:
    styled_heading(doc, "2. PLANTEAMIENTO DEL PROBLEMA", level=1)

    add_paragraph(
        doc,
        (
            "En el sector financiero, los modelos de riesgo crediticio son herramientas fundamentales "
            "para estimar las tres metricas principales del riesgo: la probabilidad de incumplimiento "
            "(PD), la perdida dado incumplimiento (LGD) y la exposicion en caso de incumplimiento "
            "(EAD) [4]. Estas metricas alimentan directamente "
            "el calculo de las perdidas esperadas (ECL = PD x LGD x EAD), que a su vez determinan las "
            "provisiones contables bajo IFRS 9 y los requerimientos de capital regulatorio bajo "
            "Basilea III [9], [6]."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Sin embargo, estos modelos presentan una limitacion recurrente y critica: suelen entregar "
            "predicciones puntuales sin una adecuada calibracion y sin cuantificar explicitamente la "
            "incertidumbre asociada. Esta deficiencia tiene un impacto multidimensional. En primer "
            "lugar, la ausencia de cuantificacion de incertidumbre puede conducir a una subestimacion "
            "sistematica del riesgo, generando perdidas inesperadas que erosionan el capital. Los "
            "modelos de machine learning, a pesar de su superior capacidad discriminativa (AUC-ROC), "
            "frecuentemente producen probabilidades mal calibradas que distorsionan las decisiones de "
            "originacion y pricing [15], [2]. En segundo "
            "lugar, ante incertidumbre no cuantificada, las entidades adoptan posturas conservadoras "
            "que resultan en exceso de capital inmovilizado y provisiones excesivas. En tercer lugar, "
            "los reguladores exigen cada vez mas que los modelos cumplan con criterios de robustez y "
            "explicabilidad que dificilmente se satisfacen sin mecanismos de cuantificacion de "
            "incertidumbre [7]."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Para el caso especifico de LGD y EAD, que son variables continuas acotadas, los metodos "
            "de Conformalized Quantile Regression (CQR) ofrecen intervalos predictivos eficientes que "
            "se adaptan localmente a la heteroscedasticidad de los datos, produciendo intervalos mas "
            "estrechos donde el modelo es mas confiable y mas amplios donde la incertidumbre es mayor "
            "[17]. Para PD, que es una variable de clasificacion binaria, las "
            "variantes como Split Conformal y Mondrian permiten obtener probabilidades acompanadas "
            "de bandas de incertidumbre con garantias de cobertura."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Pregunta de investigacion: ¿Como mejorar la calibracion y la cuantificacion de la "
            "incertidumbre en los modelos de riesgo crediticio (PD, LGD, EAD) mediante la aplicacion "
            "de tecnicas de prediccion conformal, y cual es su desempeno comparado frente a metodos "
            "tradicionales de calibracion?"
        ),
        bold=True,
        first_line_indent=1.25,
    )


# ── 3. JUSTIFICACION ─────────────────────────────────────────────────


def build_justification(doc: Document) -> None:
    styled_heading(doc, "3. JUSTIFICACION DE LA INVESTIGACION", level=1)

    add_paragraph(
        doc,
        (
            "La presente investigacion se justifica desde multiples dimensiones que abarcan el ambito "
            "academico, regulatorio y practico, respondiendo a necesidades reales tanto del sector "
            "financiero como de la comunidad cientifica en ciencia de datos aplicada."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Desde la perspectiva academica, esta investigacion contribuye a una linea de trabajo "
            "emergente que integra la cuantificacion de incertidumbre con los modelos de machine "
            "learning aplicados al riesgo financiero. Aunque la prediccion conformal ha sido "
            "ampliamente estudiada en dominios como la medicina y la vision por computador [13], [18], "
            "su aplicacion especifica al riesgo crediticio permanece "
            "relativamente inexplorada, particularmente en lo que respecta a la modelacion conjunta de "
            "PD, LGD y EAD con garantias formales de cobertura. Los trabajos de Angelopoulos y Bates "
            "[1] han popularizado el marco teorico general, pero la literatura sobre su implementacion "
            "practica en portafolios crediticios reales es escasa, lo que posiciona a esta investigacion "
            "como una contribucion original al campo."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "En el ambito regulatorio, la investigacion ofrece un marco metodologico que se alinea "
            "directamente con los requerimientos de IFRS 9 y Basilea III. Bajo IFRS 9, las entidades "
            "financieras deben estimar perdidas crediticias esperadas (ECL) incorporando informacion "
            "prospectiva y escenarios macroeconomicos, lo que requiere modelos que no solo predigan "
            "sino que cuantifiquen la incertidumbre de manera transparente. Basilea III exige que los "
            "modelos internos de riesgo sean sometidos a validacion rigurosa, incluyendo pruebas de "
            "backtesting que verifiquen la cobertura real de las estimaciones. La prediccion conformal, "
            "con sus garantias formales de cobertura, ofrece un camino natural para satisfacer estos "
            "requisitos, fortaleciendo la trazabilidad y auditabilidad de los modelos utilizados."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Desde la perspectiva practica, la adopcion de estas tecnicas permitiria a las instituciones "
            'financieras distinguir entre predicciones de alta y baja certeza. Un modelo que dice "este '
            'cliente tiene PD de 12%" es util para ranking; un modelo que dice "PD de 12% con intervalo '
            '[8%, 17%]" es util para decision. Esa distincion permite provisionar con mas precision, '
            "defender las decisiones ante comites de riesgo, y detectar anticipadamente el deterioro "
            "crediticio (SICR) a traves del ancho del intervalo conformal."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Finalmente, esta investigacion representa una innovacion para el contexto colombiano, "
            "donde la adopcion de tecnicas avanzadas de cuantificacion de incertidumbre en modelos "
            "de riesgo crediticio es aun incipiente. La implementacion exitosa de estos metodos "
            "posicionaria a la organizacion como referente en la adopcion de modelos explicables, "
            "confiables y alineados con las mejores practicas internacionales."
        ),
        first_line_indent=1.25,
    )


# ── 4. OBJETIVOS ─────────────────────────────────────────────────────


def build_objectives(doc: Document) -> None:
    styled_heading(doc, "4. OBJETIVOS", level=1)

    styled_heading(doc, "4.1 Objetivo General", level=2)
    add_paragraph(
        doc,
        (
            "Implementar y evaluar tecnicas de prediccion conformal para mejorar la calibracion y "
            "cuantificacion de la incertidumbre en modelos de riesgo crediticio (PD, LGD, EAD), "
            "utilizando el dataset publico de LendingClub como caso de estudio, con el fin de "
            "demostrar que es posible acompanar cada prediccion de riesgo con intervalos de confianza "
            "que posean garantias formales de cobertura, mejorando asi la calidad de las decisiones "
            "crediticias y la confiabilidad regulatoria de los modelos."
        ),
        first_line_indent=1.25,
    )

    styled_heading(doc, "4.2 Objetivos Especificos", level=2)

    objectives = [
        (
            "Fundamentar teoricamente la aplicacion de prediccion conformal al riesgo crediticio "
            "mediante una revision sistematica del estado del arte, identificando las variantes "
            "metodologicas mas pertinentes (Split Conformal, Mondrian, CQR) y sus ventajas frente "
            "a metodos tradicionales de calibracion, con el proposito de establecer un marco "
            "conceptual solido que guie la experimentacion."
        ),
        (
            "Construir un dataset analitico a partir de los registros publicos de LendingClub "
            "(2007-2020), con separacion temporal estricta (out-of-time) en conjuntos de "
            "entrenamiento, calibracion y test, definiendo las variables objetivo PD, LGD y EAD "
            "de manera que se garantice la ausencia de data leakage y la validez de la evaluacion "
            "fuera de muestra."
        ),
        (
            "Implementar un modelo base de PD calibrado (CatBoost + calibracion post-hoc) que "
            "sirva como fundamento para la aplicacion de prediccion conformal, evaluando su "
            "capacidad discriminativa (AUC-ROC, Gini) y su calidad probabilistica (Brier Score, "
            "ECE), para demostrar que la calibracion es un paso necesario pero insuficiente "
            "sin cuantificacion explicita de incertidumbre."
        ),
        (
            "Aplicar prediccion conformal Mondrian (por grado de riesgo) sobre PD, y variantes "
            "adaptativas sobre LGD y EAD, evaluando cobertura empirica, eficiencia de intervalos "
            "y cobertura condicional por subgrupo, para demostrar que es posible obtener garantias "
            "de cobertura que se mantienen no solo a nivel global sino en cada segmento del "
            "portafolio, lo cual es critico para la gestion operativa del riesgo."
        ),
        (
            "Evaluar el impacto de la prediccion conformal en el contexto regulatorio de IFRS 9, "
            "cuantificando como los intervalos de incertidumbre afectan la estimacion de perdidas "
            "crediticias esperadas (ECL), la clasificacion por etapas (Stages 1-3) y la "
            "sensibilidad de las provisiones ante escenarios de estres, para demostrar que la "
            "incertidumbre conformal tiene un efecto material en la lectura financiera."
        ),
        (
            "Proponer lineamientos practicos para la adopcion de Conformal Prediction en entornos "
            "bancarios reales, incluyendo criterios de monitoreo de cobertura, alertas de "
            "degradacion y escalamiento, para facilitar la transicion de una tecnica academica "
            "a una herramienta operativa de gobernanza de modelos."
        ),
    ]

    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"{i}. ")
        r.font.bold = True
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"
        r.font.color.rgb = DARK_GRAY
        r2 = p.add_run(obj)
        r2.font.size = Pt(12)
        r2.font.name = "Times New Roman"
        r2.font.color.rgb = DARK_GRAY


# ── 5. MARCO TEORICO ─────────────────────────────────────────────────


def build_theoretical_framework(doc: Document, fig_num: int) -> int:
    styled_heading(doc, "5. MARCO TEORICO Y ESTADO DEL ARTE", level=1)

    styled_heading(doc, "5.1 Modelos de Riesgo Crediticio", level=2)
    add_paragraph(
        doc,
        (
            "Los modelos de riesgo crediticio buscan estimar la probabilidad y severidad de las perdidas "
            "asociadas al incumplimiento de obligaciones financieras. Los tres componentes fundamentales "
            "son: PD, que mide la probabilidad de que un deudor incumpla en un horizonte temporal "
            "determinado; LGD, que estima la fraccion de la exposicion que se pierde tras considerar "
            "recuperaciones; y EAD, que cuantifica el monto expuesto al momento del incumplimiento "
            "[4]. Historicamente, la PD se ha modelado "
            "mediante regresion logistica, mientras que LGD y EAD han utilizado regresion beta, Tobit "
            "o regresion lineal con transformaciones. Modelos de machine learning como Gradient Boosting "
            "han demostrado mejoras significativas en discriminacion, pero frecuentemente a costa de la "
            "calibracion [12], [2]."
        ),
        first_line_indent=1.25,
    )

    styled_heading(doc, "5.2 El Problema de la Calibracion", level=2)
    add_paragraph(
        doc,
        (
            "La calibracion se refiere a la correspondencia entre las probabilidades predichas y las "
            "frecuencias observadas. Un modelo bien calibrado es aquel en el que, de todos los creditos "
            "con PD asignada del 10%, aproximadamente el 10% efectivamente incumple [15]. "
            "Los metodos tradicionales de post-hoc calibracion incluyen Platt Scaling "
            "(transformacion sigmoidal) [16] e Isotonic Regression (funcion monotona no decreciente). Sin "
            "embargo, estos metodos corrigen el nivel probabilistico pero no cuantifican la "
            'incertidumbre: un modelo calibrado dice "12% de PD" pero no indica si ese 12% es una '
            "estimacion estable o fragil. Ademas, carecen de garantias formales de cobertura y pueden "
            "ser sensibles al sobreajuste en el conjunto de calibracion [19]."
        ),
        first_line_indent=1.25,
    )

    styled_heading(doc, "5.3 Prediccion Conformal: Fundamentos Teoricos", level=2)
    add_paragraph(
        doc,
        (
            "La prediccion conformal fue introducida por Vovk, Gammerman y Shafer [18] en su trabajo "
            '"Algorithmic Learning in a Random World". El marco se fundamenta en el concepto de '
            "nonconformity scores, que miden cuan inusual es una nueva observacion respecto a un "
            "conjunto de referencia. La idea central es que, bajo el supuesto de intercambiabilidad "
            "(exchangeability) de los datos, es posible construir intervalos de prediccion que contengan "
            "el valor verdadero con probabilidad al menos (1 - alpha), sin supuestos parametricos. "
            "La garantia formal es: P(Y_nuevo ∈ C(X_nuevo)) >= 1 - alpha, donde C(X) es el conjunto "
            "de prediccion conformal [1]."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "La variante mas utilizada es Split Conformal Prediction (SCP), que divide los datos en "
            "un conjunto de entrenamiento (para ajustar el modelo) y un conjunto de calibracion (para "
            "calcular los quantiles de los nonconformity scores). Es computacionalmente eficiente y "
            "facil de implementar. Extensiones como Jackknife+ abordan la perdida de poder estadistico "
            "al no usar todos los datos [3]. Para datos heteroscedasticos, "
            "Conformalized Quantile Regression (CQR) combina regresion cuantilica con el marco "
            "conformal para producir intervalos que se adaptan localmente a la variabilidad "
            "[17]."
        ),
        first_line_indent=1.25,
    )

    styled_heading(doc, "5.4 Prediccion Conformal Mondrian", level=2)
    add_paragraph(
        doc,
        (
            "Una limitacion de Split Conformal global es que la garantia de cobertura es marginal: "
            "se cumple en promedio sobre toda la poblacion, pero puede fallar en subgrupos especificos. "
            "Por ejemplo, un portafolio con 70% de grado A y 5% de grado G puede mostrar 90% de "
            "cobertura global mientras que grado G solo alcanza 58%. La variante Mondrian resuelve "
            "este problema calculando los quantiles de no conformidad por subgrupo (por ejemplo, por "
            "grado de riesgo), garantizando cobertura condicional por particion. Esto es operativamente "
            "critico en riesgo crediticio, donde las decisiones se toman por segmento y una garantia "
            "solo promedio puede ocultar fallos en los segmentos mas riesgosos."
        ),
        first_line_indent=1.25,
    )

    fig_num = add_figure(
        doc,
        FIGURES / "figure2.png",
        "Ilustracion conceptual de la prediccion conformal: cada prediccion se acompana "
        "de un intervalo cuyo ancho refleja la incertidumbre del modelo.",
        fig_num,
        width_inches=5.0,
    )

    styled_heading(doc, "5.5 Aplicaciones de CP en el Sector Financiero", level=2)
    add_paragraph(
        doc,
        (
            "Bellotti [5] fue uno de los primeros en aplicar CP a credit scoring, demostrando que "
            "los conjuntos de prediccion conformal proporcionan informacion valiosa sobre la "
            "confiabilidad de las decisiones de credito. Javanmardi y Vovk [10] extendieron los "
            "predictores Venn-Abers para mejorar la calibracion de PD en portafolios bancarios. Sin "
            "embargo, la literatura presenta vacios significativos: la mayoria de estudios se enfocan "
            "exclusivamente en PD, dejando sin explorar la aplicacion conjunta a LGD y EAD. Ademas, "
            "pocos trabajos evaluan el impacto practico de CP en metricas regulatorias especificas "
            "como las provisiones bajo IFRS 9. Esta investigacion busca cerrar estos vacios mediante "
            "una evaluacion integral de CP en los tres componentes del riesgo crediticio."
        ),
        first_line_indent=1.25,
    )

    return fig_num


# ── 6. METODOLOGIA ───────────────────────────────────────────────────


def build_methodology(doc: Document, fig_num: int) -> int:
    styled_heading(doc, "6. METODOLOGIA", level=1)

    add_paragraph(
        doc,
        (
            "La investigacion siguio un enfoque cuantitativo y experimental, estructurado en seis "
            "fases secuenciales alineadas con el marco CRISP-DM, que abarcan desde la revision del "
            "estado del arte hasta la evaluacion del impacto regulatorio."
        ),
        first_line_indent=1.25,
    )

    styled_heading(doc, "6.1 Fase 1: Revision Bibliografica", level=2)
    add_paragraph(
        doc,
        (
            "Se realizo una revision de la literatura en Scopus, arXiv y Google Scholar, utilizando "
            'terminos como "conformal prediction credit risk", "uncertainty quantification credit '
            'scoring" y "calibration probability of default". Se incluyeron articulos publicados '
            "entre 2005 y 2026, con enfasis en contribuciones recientes (2019-2026). Se revisaron "
            "ademas guias regulatorias del Comite de Basilea y documentos tecnicos de la IFRS Foundation."
        ),
        first_line_indent=1.25,
    )

    styled_heading(doc, "6.2 Fase 2: Construccion y Preparacion del Dataset", level=2)
    add_paragraph(
        doc,
        (
            "Se utilizo el dataset publico de LendingClub (2007-2020 Q3), disponible en Kaggle [11], con "
            "aproximadamente 2.9 millones de registros y mas de 140 variables. Tras limpieza y "
            "eliminacion de variables con fuga de datos (total_pymnt, recoveries, collection_recovery_fee, "
            "entre otras), se obtuvieron 1.86 millones de prestamos con 110 columnas. Las variables "
            'objetivo se definieron asi: para PD, variable binaria basada en estados "Charged Off" y '
            '"Default"; para LGD, 1 menos la tasa de recuperacion (total_rec_prncp / funded_amnt); '
            "para EAD, la exposicion residual al momento del incumplimiento."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "El dataset se dividio en tres conjuntos con separacion temporal estricta (out-of-time, "
            "OOT) para respetar la estructura cronologica del riesgo crediticio:"
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Conjunto", "Registros", "Tasa Default", "Periodo"],
        rows=[
            ["Entrenamiento", "1,346,311", "18.52%", "2007-06 a 2017-03"],
            ["Calibracion", "237,584", "22.20%", "2017-03 a 2017-12"],
            ["Test (OOT)", "276,869", "21.98%", "2018-01 a 2020-09"],
        ],
        caption="Division temporal del dataset LendingClub.",
        table_num=1,
    )

    add_paragraph(
        doc,
        (
            "La separacion temporal es critica porque garantiza que el modelo se evalua en datos "
            "del futuro respecto al entrenamiento, simulando condiciones reales de produccion. El "
            "aumento de la tasa de default en calibracion y test (22.2% vs 18.5% en entrenamiento) "
            "refleja el deterioro macroeconomico del periodo 2017-2020, que incluye el inicio de la "
            "pandemia COVID-19. Este shift temporal pone a prueba la robustez del modelo y de los "
            "intervalos conformales en condiciones adversas, que es precisamente donde la "
            "cuantificacion de incertidumbre tiene mayor valor."
        ),
        first_line_indent=1.25,
    )

    styled_heading(doc, "6.3 Fase 3: Modelado Base y Calibracion", level=2)
    add_paragraph(
        doc,
        (
            "Se implementaron dos modelos de PD: regresion logistica como baseline interpretable y "
            "CatBoost como modelo de gradient boosting de alto rendimiento. CatBoost maneja "
            "nativamente valores faltantes y variables categoricas, evitando la necesidad de "
            "imputacion manual. La optimizacion de hiperparametros se realizo mediante Optuna con "
            "validacion temporal. El modelo final incorpora restricciones monotonicas sobre dos variables "
            "con senal economica inequivoca: ingreso anual (annual_inc, monotonia negativa: mayor ingreso "
            "implica menor PD) y ratio prestamo-ingreso (loan_to_income, monotonia positiva: mayor "
            "apalancamiento implica mayor PD). Estas restricciones garantizan coherencia economica en las "
            "predicciones sin sacrificio de discriminacion (AUC esencialmente invariante: 0.7127 con "
            "restricciones vs. 0.7128 sin restricciones), lo cual mejora la auditabilidad y aceptacion "
            "regulatoria del modelo base sobre el cual se aplica la capa conformal. Para la calibracion "
            "post-hoc, se implemento una politica de seleccion temporal multi-fold que evalua Platt "
            "Scaling [16], Isotonic Regression, Venn-Abers [19] y Beta Calibration "
            "sobre 4 folds temporales, seleccionando el metodo con mejor Brier Score que no degrade "
            "el AUC-ROC mas de 0.15 puntos porcentuales."
        ),
        first_line_indent=1.25,
    )

    styled_heading(doc, "6.4 Fase 4: Aplicacion de Conformal Prediction", level=2)
    add_paragraph(
        doc,
        (
            "Para PD, se implemento Mondrian Split Conformal Prediction utilizando la libreria "
            "MAPIE 1.3.0 [14] (SplitConformalRegressor). La PD calibrada se envuelve en un "
            "ProbabilityRegressor que transforma el problema de clasificacion binaria en regresion "
            "de probabilidades, permitiendo aplicar el framework conformal de regresion. Los quantiles "
            "de no conformidad se calculan por grado de riesgo (A-G), garantizando cobertura "
            "condicional por segmento. Para LGD, se implemento un enfoque adaptativo por grado y "
            "periodo temporal (direct_adaptive_grade_temporal), seleccionado entre cuatro variantes "
            "por cumplir todos los guardrails de cobertura, sesgo y eficiencia. Para EAD, se utilizo "
            "Split Conformal estandar. Se evaluaron niveles de confianza de 90% y 95%."
        ),
        first_line_indent=1.25,
    )

    fig_num = add_figure(
        doc,
        FIGURES / "methods_flowchart.png",
        "Arquitectura del pipeline de prediccion conformal propuesto: desde datos crudos "
        "hasta impacto regulatorio IFRS 9.",
        fig_num,
        width_inches=5.5,
    )

    styled_heading(doc, "6.5 Fase 5: Evaluacion Comparativa", level=2)
    add_paragraph(
        doc,
        (
            "La evaluacion se realizo con metricas disenadas para capturar diferentes aspectos del "
            "desempeno. Para discriminacion: AUC-ROC, Gini, KS. Para calibracion: Expected Calibration "
            "Error (ECE), Brier Score. Para cobertura conformal: cobertura empirica (proporcion de "
            "observaciones dentro del intervalo), cobertura minima por grado (Mondrian), ancho promedio "
            "de intervalos, y Winkler Score. Se realizaron pruebas estadisticas formales: test de "
            "Kupiec (cobertura incondicional) y test de Christoffersen (cobertura condicional e "
            "independencia). La evaluacion incluyo backtesting temporal sobre 35 meses y analisis "
            "de cobertura por grado x mes."
        ),
        first_line_indent=1.25,
    )

    styled_heading(doc, "6.6 Fase 6: Evaluacion de Impacto Regulatorio", level=2)
    add_paragraph(
        doc,
        (
            "Se simulo el calculo de provisiones bajo IFRS 9 utilizando las estimaciones de PD, LGD "
            "y EAD con y sin intervalos conformales. Se calcularon ECL bajo cuatro escenarios "
            "(baseline, mild stress, adverse, severe), se analizo la distribucion por etapas "
            "(Stage 1, 2, 3), y se evaluo la sensibilidad del ECL ante variaciones de PD y LGD. "
            "Se propuso el ancho del intervalo conformal (PD_high - PD_point) como senal adicional "
            "de SICR (Significant Increase in Credit Risk) para la clasificacion por etapas."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Categoria", "Herramientas"],
        rows=[
            ["ML / Modelado", "CatBoost 1.2.8, scikit-learn 1.6.1, Optuna 4.7"],
            ["Conformal", "MAPIE 1.3.0 (SplitConformalRegressor, Mondrian)"],
            ["Evaluacion", "Kupiec, Christoffersen, Winkler Score"],
            ["Datos", "pandas, numpy, pandera (validacion de esquemas)"],
            ["Visualizacion", "matplotlib, seaborn, SHAP 0.48"],
            ["Desarrollo", "Python 3.12, uv, ruff, pytest, loguru"],
        ],
        caption="Stack tecnologico utilizado en la investigacion.",
        table_num=2,
    )

    # ── 6.7 Respuesta a la Evaluacion del Anteproyecto ────────────────
    styled_heading(doc, "6.7 Respuesta a la Evaluacion del Anteproyecto", level=2)

    add_paragraph(
        doc,
        (
            "El anteproyecto fue evaluado por el docente Andres Felipe Garcia Ospina el 25 de "
            "febrero de 2026, con una calificacion final de 4.5/5.0 y concepto 'Aprobado con "
            "Ajustes'. La Tabla 3 presenta los criterios de evaluacion, las observaciones del "
            "evaluador y las acciones tomadas en esta investigacion para abordar cada punto."
        ),
        first_line_indent=1.25,
    )

    eval_table = doc.add_table(rows=7, cols=3)
    eval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(eval_table)

    eval_headers = ["Criterio (Nota)", "Observacion del Evaluador", "Como se Abordo"]
    for j, h in enumerate(eval_headers):
        cell = eval_table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, TABLE_HEADER_BG)

    eval_rows = [
        [
            "1. Pertinencia (4.8)",
            "La propuesta es pertinente en cuanto a la utilidad en el entorno real "
            "y esta alineada con las lineas del programa.",
            "Se mantuvo el enfoque original. Los resultados confirman la pertinencia: "
            "coberturas conformales >90% y cuantificacion de incertidumbre de $814M "
            "en provisiones ECL demuestran impacto real.",
        ],
        [
            "2. Planteamiento (4.3)",
            "El problema esta bien argumentado frente a las deficiencias de las "
            "predicciones puntuales. Sin embargo, la formulacion de los objetivos "
            "es susceptible de mejora en funcion del alcance realista del proyecto.",
            "Se reformularon los 6 objetivos especificos con estructura 'para que / "
            "porque / como lo logramos', vinculando cada uno con impacto medible. "
            "Se acoto el alcance a CP sobre PD, LGD, EAD e IFRS 9, excluyendo "
            "optimizacion robusta e inferencia causal (reservados para maestria).",
        ],
        [
            "3. Metodologia (4.6)",
            "El diseno de 6 fases es adecuado e incluye validacion de modelos base "
            "frente a calibradores. Sin embargo, el alcance es muy amplio y por ende "
            "riesgoso metodologicamente para un solo investigador en el tiempo propuesto.",
            "Se mantuvo la metodologia de 6 fases pero se focalizo el alcance: se "
            "trabajo exclusivamente con CatBoost (no XGBoost ni LightGBM comparativo), "
            "se selecciono automaticamente el calibrador optimo (Venn-Abers gano), y se "
            "concentro el analisis conformal en Mondrian como variante principal.",
        ],
        [
            "4. Innovacion (4.7)",
            "La propuesta si bien no es 100% original, es novedosa en los elementos "
            "que pretende integrar para abordar el problema de estudio.",
            "La innovacion se materializo en tres contribuciones: (1) extension de CP "
            "a la triada completa PD-LGD-EAD (la literatura solo cubre PD), (2) "
            "demostracion de +30pp en cobertura minima con Mondrian vs Global, y (3) "
            "propuesta del ancho conformal como senal de SICR para IFRS 9.",
        ],
        [
            "5. Viabilidad (3.5)",
            "El cronograma propuesto es muy ajustado para la cantidad de modelos a "
            "entrenar, calibrar y someter a simulaciones, lo cual pone en riesgo la "
            "culminacion oportuna del proyecto.",
            "Se redujo el alcance siguiendo la recomendacion: en lugar de comparar "
            "multiples algoritmos (XGBoost, LightGBM, CatBoost), se selecciono "
            "CatBoost como modelo unico por su manejo nativo de NaN y categoricas. "
            "Se automatizo la seleccion de calibrador y se priorizo profundidad "
            "(4 variantes LGD, backtesting 35 meses) sobre amplitud.",
        ],
        [
            "6. Comunicacion (4.8)",
            "La redaccion y los elementos generales de forma son adecuados.",
            "Se mantuvo el estandar de redaccion. Se agrego rigor en la presentacion "
            "de resultados con tablas comparativas, pruebas estadisticas formales "
            "(Kupiec, Christoffersen) y visualizaciones detalladas por grado.",
        ],
    ]

    for i, row_data in enumerate(eval_rows):
        for j, val in enumerate(row_data):
            cell = eval_table.rows[i + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"
            if i % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run("Tabla 3. Respuesta a los criterios de evaluacion del anteproyecto.")
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    r.italic = True
    r.font.color.rgb = DARK_GRAY

    # ── 6.8 Cambios Metodologicos Relevantes ──────────────────────────
    styled_heading(doc, "6.8 Cambios Metodologicos respecto al Anteproyecto", level=2)

    add_paragraph(
        doc,
        (
            "Durante la ejecucion del proyecto, se realizaron ajustes metodologicos informados "
            "por los resultados experimentales y la retroalimentacion de la evaluacion. La "
            "Tabla 4 resume los cambios principales y su justificacion."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Aspecto", "Anteproyecto", "Implementacion Final", "Razon del Cambio"],
        rows=[
            [
                "Modelo base PD",
                "XGBoost como modelo principal",
                "CatBoost 1.2.8",
                "CatBoost maneja NaN y categoricas nativamente, "
                "evitando pasos de imputacion y encoding que "
                "introducen decisiones arbitrarias. Reduce riesgo "
                "de data leakage por transformaciones.",
            ],
            [
                "Alcance de modelos",
                "Comparativa de multiples algoritmos (LR, XGBoost, LightGBM, CatBoost)",
                "LR (baseline) + CatBoost (campeon)",
                "Siguiendo recomendacion del evaluador (viabilidad), "
                "se priorizo profundidad en CP sobre amplitud "
                "comparativa de algoritmos.",
            ],
            [
                "Seleccion calibrador",
                "Evaluacion manual de Platt e Isotonic",
                "Politica automatica multi-fold temporal",
                "La seleccion automatica con 4 folds temporales "
                "y criterio Brier + AUC-guard elimina sesgo del "
                "investigador. Venn-Abers gano consistentemente.",
            ],
            [
                "Variantes LGD",
                "Split Conformal unico",
                "Benchmark de 4 variantes con guardrails",
                "LGD tiene heteroscedasticidad por grado; solo la "
                "variante adaptive_grade_temporal paso todos los "
                "guardrails de cobertura y eficiencia.",
            ],
            [
                "Conformal PD",
                "Split Conformal global",
                "Mondrian por grado (A-G)",
                "Split global mostro 58.69% min cobertura por grado, "
                "inaceptable operativamente. Mondrian logro 87.82% "
                "(+29pp) con intervalos 21% mas eficientes.",
            ],
            [
                "Libreria conformal",
                "MAPIE (version no especificada)",
                "MAPIE 1.3.0 (SplitConformalRegressor)",
                "Migracion a API v1.3.0 con SplitConformalRegressor "
                "y ProbabilityRegressor wrapper, siguiendo "
                "la documentacion oficial actualizada.",
            ],
        ],
        caption="Cambios metodologicos respecto al anteproyecto y su justificacion.",
        table_num=4,
    )

    return fig_num


# ── 7. CONSIDERACIONES ETICAS ────────────────────────────────────────


def build_ethics(doc: Document) -> None:
    styled_heading(doc, "7. CONSIDERACIONES ETICAS", level=1)

    add_paragraph(
        doc,
        (
            "La presente investigacion se adhiere a los mas altos estandares de etica en la "
            "investigacion cientifica. En cuanto a la privacidad y proteccion de datos, el estudio "
            "utiliza exclusivamente datos de acceso publico provenientes de la plataforma LendingClub, "
            "previamente anonimizados y sin informacion de identificacion personal (PII). Se adoptaron "
            "medidas adicionales para garantizar que ningun analisis individual pueda conducir a la "
            "re-identificacion de deudores, siguiendo las directrices de la Ley de Proteccion de "
            "Datos Personales de Colombia (Ley 1581 de 2012)."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "En terminos de transparencia y reproducibilidad, todo el codigo fuente, los notebooks "
            "de experimentacion y los resultados intermedios fueron documentados para permitir la "
            "reproduccion completa por terceros. Se reportaron tanto los resultados positivos como "
            "las limitaciones, incluyendo alertas de cobertura que no pasaron todos los checks "
            "estadisticos. La integridad academica se garantiza mediante el uso riguroso de "
            "referencias bibliograficas y el cumplimiento de las normas de citacion."
        ),
        first_line_indent=1.25,
    )


# ── 8. RESULTADOS ────────────────────────────────────────────────────


def build_results(doc: Document, fig_num: int) -> int:
    styled_heading(doc, "8. RESULTADOS", level=1)

    add_paragraph(
        doc,
        (
            "A continuacion se presentan los resultados obtenidos, organizados desde el analisis "
            "exploratorio de datos hasta el impacto regulatorio en provisiones IFRS 9. Todos los "
            "resultados corresponden a evaluaciones sobre el conjunto de test out-of-time (276,869 "
            "prestamos, periodo 2018-01 a 2020-09), a menos que se indique lo contrario."
        ),
        first_line_indent=1.25,
    )

    # ── 8.1 EDA ──────────────────────────────────────────────────────
    styled_heading(doc, "8.1 Analisis Exploratorio de Datos", level=2)

    add_paragraph(
        doc,
        (
            "El dataset de LendingClub contiene prestamos originados entre 2007 y 2020, con un "
            "gradiente de riesgo claro por grado de riesgo asignado por la plataforma. La Tabla 5 "
            "muestra la distribucion de prestamos y la tasa de default por grado en el conjunto "
            "de entrenamiento, confirmando la senal economica del riesgo."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Grado", "Prestamos", "Tasa Default", "Tasa Interes Prom."],
        rows=[
            ["A", "229,809", "5.63%", "7.2%"],
            ["B", "401,151", "12.29%", "10.9%"],
            ["C", "383,391", "20.64%", "14.0%"],
            ["D", "198,090", "28.31%", "17.7%"],
            ["E", "94,233", "36.47%", "21.0%"],
            ["F", "31,748", "43.44%", "24.5%"],
            ["G", "7,889", "47.71%", "27.0%"],
        ],
        caption="Distribucion de prestamos y tasa de default por grado de riesgo (conjunto de entrenamiento).",
        table_num=5,
    )

    add_paragraph(
        doc,
        (
            "El gradiente de default escala monotonicamente de A (5.63%) a G (47.71%), lo que confirma "
            "que el grado de riesgo captura una senal economica real y no solo volumen. La relacion "
            "entre grado y tasa de interes promedio revela el mecanismo de pricing por riesgo de la "
            "plataforma: los grados mas riesgosos (F, G) pagan tasas 3.5 veces superiores a los "
            "grados A. Esta estructura de riesgo hace que el grado sea una variable de particion natural "
            "para Mondrian Conformal Prediction, ya que cada segmento tiene un perfil de riesgo "
            "significativamente diferente que justifica quantiles de no conformidad independientes."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "En terminos de composicion, el 74.6% de los prestamos son a 36 meses (1.00M) y "
            "el 25.4% a 60 meses (340K). Los grados A y B concentran el 46.9% del volumen, mientras "
            "que F y G representan solo el 2.9%, lo que implica que el enfoque global de conformal "
            "estara dominado por los grados con mas volumen, potencialmente subatendiendo los grados "
            "minoritarios. Las variables con mayor proporcion de valores faltantes "
            "incluyen mths_since_last_delinq (51.6%) y mths_since_last_record (84.3%), patron tipico "
            "en datos crediticios donde la ausencia indica que el evento no ha ocurrido."
        ),
        first_line_indent=1.25,
    )

    fig_num = add_figure(
        doc,
        NB_IMAGES / "01_eda_lending_club" / "default_rate_by_grade.png",
        "Tasa de default por grado de riesgo en el conjunto de entrenamiento. "
        "El gradiente monotonicamente creciente valida la senal economica del grado.",
        fig_num,
        width_inches=4.5,
    )

    fig_num = add_figure(
        doc,
        NB_IMAGES / "01_eda_lending_club" / "correlation_matrix.png",
        "Matriz de correlacion de las principales variables numericas. Se observa alta "
        "correlacion entre tasa de interes y grado (r > 0.9), confirmando que ambas "
        "capturan la misma senal de riesgo.",
        fig_num,
        width_inches=4.5,
    )

    # ── 8.2 PD ───────────────────────────────────────────────────────
    styled_heading(doc, "8.2 Desempeno del Modelo de PD", level=2)

    add_paragraph(
        doc,
        (
            "Se entrenaron y evaluaron dos modelos de PD sobre el conjunto de test OOT. "
            "La Tabla 6 resume las metricas comparativas. Las metricas ECE y KS se reportan "
            "unicamente para el modelo calibrado final, ya que son las metricas de calidad "
            "probabilistica relevantes para la toma de decisiones post-calibracion; para los modelos "
            "sin calibrar, el AUC-ROC y Brier Score capturan la capacidad discriminativa y la calidad "
            "de las predicciones brutas."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Modelo", "AUC-ROC", "Gini", "Brier", "D2-Brier", "ECE", "KS"],
        rows=[
            ["Reg. Logistica", "0.679", "0.358", "0.233", "-0.356", "—", "—"],
            ["CatBoost (default)", "0.713", "0.426", "0.210", "-0.222", "—", "—"],
            ["CatBoost (tuned)", "0.713", "0.426", "0.210", "-0.222", "—", "—"],
            ["CatBoost (calibrado)", "0.713", "0.425", "0.155", "0.099", "0.007", "0.313"],
        ],
        caption="Comparacion de modelos PD en el conjunto de test OOT (276,869 prestamos).",
        table_num=6,
        footnote=(
            "Nota: ECE y KS se evaluan solo para el modelo calibrado final. D2-Brier negativo "
            "indica peor calibracion que el modelo trivial (prediccion constante de la prevalencia). "
            "Los guiones (—) indican metricas no evaluadas para esa configuracion."
        ),
    )

    add_paragraph(
        doc,
        (
            "El modelo campeon CatBoost calibrado con Venn-Abers alcanzo un AUC-ROC de "
            "0.7127, un Brier Score de 0.1546 y un ECE de 0.0067. La calibracion Venn-Abers mejoro "
            "sustancialmente la calidad probabilistica: el Brier Score paso de 0.210 a 0.155, y el "
            "D2-Brier Score paso de -0.222 (peor que el modelo trivial) a +0.099 (mejor que el "
            "trivial), todo sin degradar la discriminacion (el AUC cae solo 0.0001, de 0.7128 a "
            "0.7127). Un D2-Brier negativo significa que el modelo sin calibrar tiene peor calidad "
            "probabilistica que simplemente predecir la tasa de default promedio para todos los "
            "prestamos, lo cual subraya la importancia critica de la calibracion."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Un aspecto relevante del modelo campeon es la incorporacion de restricciones monotonicas "
            "en dos variables con senal economica inequivoca: ingreso anual (annual_inc, monotonia "
            "negativa) y ratio prestamo-ingreso (loan_to_income, monotonia positiva). La auditoria de "
            "monotonicity confirmo cero disrupciones y 0% de tasa de violacion en ambas variables. "
            "Este resultado es significativo para la prediccion conformal porque garantiza que el modelo "
            "base produce predicciones economicamente coherentes: si un prestatario tiene mayor ingreso, "
            "su PD predicha sera menor o igual, y si su apalancamiento es mayor, su PD sera mayor o "
            "igual. Los intervalos conformales construidos sobre este modelo heredan esa coherencia, "
            "lo cual mejora la interpretabilidad y aceptacion regulatoria de las bandas de incertidumbre."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Sin embargo, incluso el modelo calibrado con ECE de 0.0067 aun no responde a la "
            "pregunta fundamental para la gestion del riesgo: ¿con que confianza puedo tomar una "
            "decision basada en esta PD? Un prestamo con PD calibrada de 15% podria ser una estimacion "
            "estable (intervalo [13%, 17%]) o fragil (intervalo [5%, 28%]). Esa distincion es "
            "precisamente lo que resuelve la prediccion conformal y lo que se aborda en las siguientes "
            "secciones."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Las cinco variables mas importantes identificadas por SHAP fueron: tasa de interes "
            "(int_rate), plazo del prestamo (term), puntaje FICO (fico_score), tipo de vivienda "
            "(home_ownership) y razon deuda-ingreso (dti). Estas variables son consistentes con "
            "la literatura de riesgo crediticio y refuerzan la interpretabilidad del modelo, lo "
            "que es relevante porque la prediccion conformal hereda la estructura del modelo base: "
            "si el modelo discrimina bien, los intervalos conformales seran mas informativos."
        ),
        first_line_indent=1.25,
    )

    fig_num = add_figure(
        doc,
        NB_IMAGES / "03_pd_modeling" / "roc_curves.png",
        "Curvas ROC comparativas de los modelos PD en el conjunto OOT. CatBoost "
        "supera al baseline logistico en +0.029 AUC, con ganancia concentrada en "
        "los segmentos de alto riesgo (grados D-G).",
        fig_num,
        width_inches=4.5,
    )

    fig_num = add_figure(
        doc,
        NB_IMAGES / "03_pd_modeling" / "calibration_curves.png",
        "Curvas de calibracion: CatBoost sin calibrar (izquierda, desviacion visible "
        "de la diagonal) vs. calibrado con Venn-Abers (derecha, alineado "
        "con la diagonal perfecta). La calibracion corrige la subestimacion sistematica.",
        fig_num,
        width_inches=4.5,
    )

    fig_num = add_figure(
        doc,
        NB_IMAGES / "03_pd_modeling" / "feature_importance.png",
        "Importancia global de features por SHAP values (top 20). La tasa de interes "
        "domina, seguida por el plazo y el puntaje FICO, variables alineadas con la "
        "teoria de riesgo crediticio.",
        fig_num,
        width_inches=4.5,
    )

    # ── 8.3 Conformal PD Mondrian ────────────────────────────────────
    styled_heading(doc, "8.3 Intervalos Conformales para PD (Mondrian)", level=2)

    add_paragraph(
        doc,
        (
            "Se aplico Mondrian Split Conformal Prediction sobre la PD calibrada, calculando los "
            "quantiles de no conformidad por grado de riesgo (A-G). La Tabla 7 resume las metricas "
            "de cobertura y eficiencia para los niveles de confianza del 90% y 95%."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Metrica", "Nivel 90%", "Nivel 95%"],
        rows=[
            ["Cobertura empirica global", "92.42%", "95.93%"],
            ["Cobertura minima por grado", "89.01%", "93.4%*"],
            ["Ancho medio de intervalo", "0.7546", "0.87*"],
            ["Winkler Score", "1.209", "1.152"],
            ["Tasa de violacion", "7.58%", "4.07%"],
            ["Kupiec p-value", "< 0.001", "< 0.001"],
            ["Christoffersen p_ind", "0.512", "0.034"],
        ],
        caption="Metricas de cobertura y eficiencia de intervalos conformales PD (Mondrian).",
        table_num=7,
        footnote=(
            "* Valores estimados a partir de los datos de backtesting. La tasa de violacion "
            "es el complemento de la cobertura (1 - cobertura). Checks del gate de politica: "
            "7/13 pasaron; 5 alertas activas (todas de tipo warning, ninguna critica)."
        ),
    )

    add_paragraph(
        doc,
        (
            "La cobertura empirica al nivel del 90% alcanzo 92.42%, superando el objetivo nominal "
            "en 2.42 puntos porcentuales. Esto significa que de los 276,869 prestamos del test OOT, "
            "en 255,890 (92.42%) el valor real de default quedo dentro del intervalo conformal "
            "predicho, mientras que 20,979 (7.58%) quedaron fuera. La ligera sobrecobertura indica "
            "que los intervalos son conservadores, lo cual es deseable en un contexto regulatorio "
            "donde subestimar la incertidumbre tiene consecuencias mas graves que sobreestimarla."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "La cobertura minima por grado fue 89.01%, demostrando que el enfoque Mondrian logra "
            "coberturas aceptables incluso en los segmentos mas desafiantes. En terminos practicos, "
            "esto significa que ni siquiera el peor grado (con menor volumen de datos y mayor "
            "variabilidad) cae por debajo del 89%, lo cual seria inaceptable con el enfoque global "
            "donde se observo un minimo de 58.69% (seccion 8.4)."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Las pruebas de Kupiec rechazaron la cobertura exacta al 90% (p-value < 0.001), "
            "lo cual es esperado con 276,869 observaciones: con muestras tan grandes, incluso "
            "desviaciones de 1.76pp del nominal son estadisticamente significativas. Lo relevante "
            "operativamente es que la cobertura supera el nivel objetivo, validando la utilidad "
            "practica. El test de Christoffersen confirmo independencia de las violaciones al 90% "
            "(p_ind = 0.512 > 0.05), indicando que las violaciones no se agrupan en periodos "
            "especificos, sino que se distribuyen aleatoriamente en el tiempo. Al 95%, la "
            "independencia se rechaza marginalmente (p_ind = 0.034), sugiriendo leve agrupamiento "
            "que merece monitoreo."
        ),
        first_line_indent=1.25,
    )

    fig_num = add_figure(
        doc,
        NB_IMAGES / "04_conformal_prediction" / "coverage_by_grade.png",
        "Cobertura empirica por grado de riesgo (Mondrian Conformal Prediction). "
        "Todos los grados superan el 89%, con los grados de mayor volumen (A-C) "
        "mas cercanos al nominal del 90%.",
        fig_num,
        width_inches=4.5,
    )

    fig_num = add_figure(
        doc,
        NB_IMAGES / "04_conformal_prediction" / "interval_width_distribution.png",
        "Distribucion del ancho de intervalos conformales PD al nivel del 90%. "
        "La distribucion bimodal refleja la diferencia entre grados de bajo riesgo "
        "(intervalos estrechos) y alto riesgo (intervalos amplios).",
        fig_num,
        width_inches=4.5,
    )

    # ── 8.4 Mondrian vs Global ───────────────────────────────────────
    styled_heading(doc, "8.4 Comparacion: Mondrian vs. Split Global", level=2)

    add_paragraph(
        doc,
        (
            "Para demostrar la ventaja operativa de Mondrian sobre Split Conformal global, se "
            "compararon ambas variantes sobre el mismo conjunto de test. La Tabla 8 resume los "
            "resultados comparativos."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Metrica", "Global", "Mondrian", "Diferencia"],
        rows=[
            ["Cobertura 90%", "89.84%", "92.42%", "+2.58 pp"],
            ["Min. cobertura por grado", "58.69%", "89.01%", "+30.32 pp"],
            ["Ancho medio 90%", "0.955", "0.7546", "-21.0%"],
        ],
        caption="Comparacion entre Split Conformal global y Mondrian por grado.",
        table_num=8,
    )

    add_paragraph(
        doc,
        (
            "La diferencia mas relevante es en la cobertura minima por grado: Global alcanza solo "
            "58.69% en su peor subgrupo, mientras que Mondrian sube a 89.01% (+30 puntos "
            "porcentuales). Para poner esto en contexto operativo: con el enfoque global, un banco "
            "que tiene 7,889 prestamos de grado G veria que en el 41.3% de esos prestamos el "
            "intervalo conformal no contiene el valor real. Esto invalida por completo la garantia "
            "de cobertura para ese segmento y hace que los intervalos sean inservibles para la "
            "toma de decisiones en grados de alto riesgo."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Ademas, Mondrian logra intervalos 21.0% mas eficientes (ancho "
            "promedio de 0.7546 vs 0.955). Esto parece contradictorio —mejor cobertura con "
            "intervalos mas estrechos— pero se explica porque Mondrian calibra los quantiles "
            "por grado: los grados A y B, que son los mas predecibles, obtienen intervalos estrechos "
            "y precisos, mientras que los grados F y G obtienen intervalos mas amplios pero "
            "que reflejan fielmente su mayor incertidumbre. Global, en cambio, usa un unico "
            "quantil que sobredimensiona los grados faciles y subdimensiona los dificiles."
        ),
        first_line_indent=1.25,
    )

    fig_num = add_figure(
        doc,
        NB_IMAGES / "04_conformal_prediction" / "coverage_width_tradeoff.png",
        "Trade-off entre cobertura y ancho de intervalos conformales. Mondrian "
        "domina a Global en ambas dimensiones: mejor cobertura condicional con "
        "intervalos mas eficientes.",
        fig_num,
        width_inches=4.5,
    )

    # ── 8.5 LGD / EAD ───────────────────────────────────────────────
    styled_heading(doc, "8.5 Intervalos Conformales para LGD y EAD", level=2)

    add_paragraph(
        doc,
        (
            "Los intervalos conformales se extendieron a LGD y EAD sobre el subconjunto de defaults "
            "(60,850 prestamos en test OOT). Para LGD, se evaluaron cuatro variantes conformales "
            "con un benchmark sistematico que incluye guardrails de cobertura, sesgo y eficiencia. "
            "La Tabla 9 compara las variantes."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=[
            "Variante LGD",
            "Cob. 90%",
            "Cob. 95%",
            "Min Grado 90%",
            "Ancho 90%",
            "Guardrails",
        ],
        rows=[
            ["Two-stage split", "78.19%", "87.30%", "53.56%", "0.568", "No pasa"],
            ["Direct split", "78.64%", "87.64%", "54.12%", "0.568", "No pasa"],
            ["Direct CQR", "74.52%", "84.47%", "70.51%", "0.540", "No pasa"],
            ["Adaptive grade-temporal", "90.50%", "95.50%", "90.47%", "0.496", "Pasa todos"],
        ],
        caption="Benchmark de variantes conformales para LGD (60,850 defaults en test OOT).",
        table_num=9,
    )

    add_paragraph(
        doc,
        (
            "Solo la variante direct_adaptive_grade_temporal paso todos los guardrails. Las tres "
            "variantes rechazadas muestran un patron comun: subcoberturas severas (74-78% vs el "
            "objetivo de 90%), especialmente en las coberturas por grado (53-70% en el peor grado). "
            "Esto demuestra que LGD, al ser una variable continua acotada en [0, 1] con alta "
            "heteroscedasticidad por grado, requiere un enfoque conformal que se adapte a la "
            "estructura del portafolio. La variante ganadora ajusta los quantiles conformales de "
            "manera online por grado y periodo temporal, logrando no solo la cobertura mas alta "
            "(90.50%) sino tambien el ancho mas eficiente (0.496, un 13% menor que la referencia)."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Para EAD, el modelo alcanzo cobertura del "
            "91.20% al 90% y 95.28% al 95%, con un R-cuadrado de 0.9999 en la prediccion puntual. "
            "La alta precision de EAD se explica porque la exposicion al default esta fuertemente "
            "determinada por el saldo vivo del prestamo, que es una variable contractual con baja "
            "incertidumbre. Los intervalos conformales para EAD son, en consecuencia, estrechos "
            "(ancho medio de $132.58 USD al 90%), reflejando correctamente la baja incertidumbre."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Este resultado es significativo porque demuestra que la prediccion conformal es "
            "aplicable a la triada completa PD-LGD-EAD, no solo a PD. La extension a LGD y EAD "
            "cierra un vacio en la literatura existente, donde trabajos como Bellotti [5] se enfocan "
            "exclusivamente en PD, y habilita el calculo de intervalos de "
            "confianza para la perdida esperada completa (ECL = PD x LGD x EAD)."
        ),
        first_line_indent=1.25,
    )

    # ── 8.6 IFRS9 ───────────────────────────────────────────────────
    styled_heading(doc, "8.6 Impacto Regulatorio: IFRS 9 y ECL", level=2)

    add_paragraph(
        doc,
        (
            "Se evaluo el impacto de los intervalos conformales en el calculo de perdidas crediticias "
            "esperadas (ECL) bajo IFRS 9 [9]. La Tabla 10 muestra el ECL bajo cuatro escenarios "
            "que van desde el baseline (PD puntual) hasta el escenario severo (PD en el limite "
            "superior del intervalo conformal al 95%)."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Escenario", "ECL (USD)", "Uplift vs Baseline", "PD utilizada"],
        rows=[
            ["Baseline (PD puntual)", "$1,001M", "—", "PD calibrada"],
            ["Mild stress", "$1,241M", "+24.0%", "PD + 0.5 * ancho CP"],
            ["Adverse", "$1,501M", "+50.0%", "PD + 0.75 * ancho CP"],
            ["Severe (PD alta conformal)", "$1,799M", "+79.7%", "PD_high (95%)"],
        ],
        caption="Estimacion de ECL bajo escenarios de estres con intervalos conformales.",
        table_num=10,
    )

    add_paragraph(
        doc,
        (
            "El ECL baseline (usando PD puntual calibrada) es de $1,001M. Cuando se utiliza el limite "
            "superior del intervalo conformal al 95% (escenario severe), el ECL sube a $1,799M, un "
            "incremento del 79.7%. Este rango de $798M entre el escenario base y el severo cuantifica "
            "de manera explicita la incertidumbre en las provisiones. Sin prediccion conformal, este "
            "rango simplemente no existe: el banco solo ve un numero ($1,001M) sin saber si es una "
            "estimacion robusta o fragil."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "La distribucion por etapas IFRS 9 fue: Stage 1 (sin deterioro significativo) 34.51% "
            "(95,547 prestamos), Stage 2 (SICR detectado) 43.01% (119,071 prestamos), y Stage 3 "
            "(impago) 22.48% (62,251 prestamos). La alta proporcion de Stage 2 refleja que el "
            "periodo 2018-2020 incluye el inicio de la pandemia, cuando se activo un incremento "
            "significativo de riesgo crediticio a nivel del portafolio completo."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Un hallazgo relevante de esta investigacion es que "
            "el ancho del intervalo conformal (PD_high - PD_point) puede utilizarse como senal "
            "adicional de SICR para la clasificacion por etapas: un prestamo cuyo intervalo se "
            "amplifica significativamente entre periodos indica deterioro, incluso si la PD puntual "
            "no ha cambiado lo suficiente para activar los umbrales tradicionales. Esta propuesta "
            "complementa los criterios cuantitativos existentes (incremento absoluto o relativo de PD) "
            "con una dimension de incertidumbre que captura la fragilidad de la estimacion."
        ),
        first_line_indent=1.25,
    )

    fig_num = add_figure(
        doc,
        NB_IMAGES / "05_time_series_forecasting" / "ifrs9_scenario_fan_chart.png",
        "Fan chart de escenarios IFRS 9 con intervalos de confianza. La franja "
        "entre baseline ($1,001M) y severe ($1,799M) cuantifica la incertidumbre "
        "material en las provisiones.",
        fig_num,
        width_inches=4.5,
    )

    return fig_num


# ── 9. DISCUSION ─────────────────────────────────────────────────────


def build_discussion(doc: Document) -> None:
    styled_heading(doc, "9. DISCUSION DE RESULTADOS", level=1)

    styled_heading(doc, "9.1 Hallazgos Principales", level=2)

    add_paragraph(
        doc,
        (
            "Los resultados obtenidos permiten responder afirmativamente la pregunta de investigacion: "
            "las tecnicas de prediccion conformal mejoran efectivamente la cuantificacion de "
            "incertidumbre en modelos de riesgo crediticio, con impacto demostrable en la lectura "
            "regulatoria."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "El primer hallazgo clave es la complementariedad entre calibracion y prediccion conformal. "
            "La calibracion Venn-Abers redujo el ECE a 0.0067, corrigiendo el nivel probabilistico del "
            "modelo. La prediccion conformal, sobre esa PD ya calibrada, agrego una banda de "
            "incertidumbre con cobertura controlada. Estos dos mecanismos no compiten: la calibracion "
            'responde "¿cuanto riesgo hay?", y la prediccion conformal responde "¿con que confianza '
            'lo digo?". Esta complementariedad es consistente con el trabajo de Javanmardi y Vovk [10], '
            "donde Venn-Abers se utiliza como calibrador multiprobabilistico; en nuestro pipeline, la "
            "combinacion secuencial CatBoost + Venn-Abers + Mondrian fue altamente efectiva."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "El segundo hallazgo es la superioridad operativa de Mondrian sobre Split Conformal global. "
            "La diferencia de +30 puntos porcentuales en cobertura minima por grado (89.01% vs 58.69%) "
            "no es solo estadisticamente significativa sino operativamente critica: en un banco, las "
            "decisiones de credito se toman por segmento, y una garantia solo promedio puede ocultar "
            "fallos graves en los segmentos de mayor riesgo. Ademas, Mondrian logra intervalos 21% "
            "mas eficientes, demostrando que no hay trade-off entre cobertura condicional y eficiencia."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "El tercer hallazgo es la extension exitosa a la triada PD-LGD-EAD. La mayoria de la "
            "literatura existente [5] se enfoca exclusivamente en PD. Este trabajo "
            "demuestra que la prediccion conformal es aplicable tambien a LGD (90.50% de cobertura "
            "con la variante adaptativa) y EAD (91.20%). Para LGD, la variante "
            "direct_adaptive_grade_temporal fue la unica de cuatro que paso todos los guardrails, "
            "subrayando que el framework conformal debe adaptarse a las caracteristicas de cada "
            "componente y no aplicarse de manera generica."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "El cuarto hallazgo es el impacto material en la lectura regulatoria IFRS 9. El rango "
            "de ECL entre $1,001M y $1,799M (+79.7%) proporciona a los comites de riesgo una "
            "cuantificacion explicita de la incertidumbre en las provisiones. El uso del ancho del "
            "intervalo conformal como senal de SICR es una innovacion metodologica que podria mejorar "
            "la deteccion temprana de deterioro crediticio."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Es importante reconocer las limitaciones. El supuesto de intercambiabilidad puede verse "
            "comprometido en periodos de crisis economica. Los intervalos para grados F y G tienden "
            "a ser mas amplios, reduciendo su utilidad practica. Las pruebas de Kupiec/Christoffersen "
            "rechazan la cobertura exacta, aunque esto se debe al tamano muestral extremadamente "
            "grande y no a una deficiencia practica. Finalmente, el gate estricto (7/13 checks) "
            "indica que quedan alertas abiertas para monitoreo operativo, lo que es esperado en un "
            "sistema de gobernanza que prioriza transparencia sobre aprobacion automatica."
        ),
        first_line_indent=1.25,
    )

    # ── 9.2 Resultados en el Marco CRISP-DM ─────────────────────────
    styled_heading(doc, "9.2 Resultados en el Marco CRISP-DM", level=2)

    add_paragraph(
        doc,
        (
            "La investigacion se estructuro siguiendo el ciclo CRISP-DM (Cross-Industry Standard "
            "Process for Data Mining). La Tabla 11 mapea cada fase del proceso con las actividades "
            "realizadas y los entregables concretos obtenidos."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Fase CRISP-DM", "Actividades Realizadas", "Entregables"],
        rows=[
            [
                "1. Comprension del Negocio",
                "Revision de normativa IFRS 9 y Basilea III. "
                "Identificacion de la brecha: predicciones puntuales "
                "sin cuantificacion de incertidumbre. Definicion de "
                "la pregunta de investigacion.",
                "Propuesta de investigacion aprobada. "
                "Pregunta de investigacion formulada. "
                "6 objetivos especificos definidos.",
            ],
            [
                "2. Comprension de los Datos",
                "EDA del dataset LendingClub (2.9M registros, 140+ "
                "variables). Analisis de distribuciones por grado, "
                "patrones de missing values, correlaciones. "
                "Identificacion de variables con data leakage.",
                "Gradiente monotono de default A(5.6%)-G(47.7%). "
                "Lista de 22 variables con leakage eliminadas. "
                "Confirmacion del grado como variable de particion.",
            ],
            [
                "3. Preparacion de los Datos",
                "Limpieza (eliminacion leakage, filtro de estados "
                "resueltos). Feature engineering (42 features). "
                "Split temporal OOT: train/cal/test.",
                "3 conjuntos OOT: 1.35M / 237K / 277K registros. "
                "3 datasets analiticos: loan_master, time_series, "
                "ead_dataset. Contrato de features (JSON).",
            ],
            [
                "4. Modelado",
                "LR baseline + CatBoost (default/tuned/monotónico). "
                "Calibracion post-hoc (Platt, Isotonic, Venn-Abers, Beta) "
                "con seleccion automatica multi-fold. Mondrian CP "
                "para PD. 4 variantes CP para LGD. Split CP para EAD.",
                "Modelo campeon: CatBoost + Venn-Abers (AUC 0.713, "
                "ECE 0.007). Intervalos conformales PD (92.42% cob.), "
                "LGD (90.50% cob.), EAD (91.20% cob.).",
            ],
            [
                "5. Evaluacion",
                "Evaluacion OOT con metricas de discriminacion, "
                "calibracion y cobertura. Pruebas de Kupiec y "
                "Christoffersen. Backtesting 35 meses. "
                "Benchmark Mondrian vs Global. Gate de politica.",
                "Mondrian +30pp vs Global en min cobertura. "
                "Independencia de violaciones confirmada (p=0.512). "
                "Gate: 7/13 checks (0 criticos). "
                "Impacto ECL: $1,001M-$1,799M (+79.7%).",
            ],
            [
                "6. Despliegue",
                "Dashboard Streamlit con 27 paginas de resultados. "
                "Pipeline reproducible con DVC (26 stages). "
                "Lineamientos de monitoreo y recalibracion.",
                "Dashboard interactivo de resultados. "
                "Pipeline end-to-end automatizado. "
                "Framework de gobernanza con alertas por grado.",
            ],
        ],
        caption="Mapeo de fases CRISP-DM con actividades y entregables del proyecto.",
        table_num=11,
    )

    # ── 9.3 Cumplimiento de Objetivos ────────────────────────────────
    styled_heading(doc, "9.3 Cumplimiento de Objetivos", level=2)

    add_paragraph(
        doc,
        (
            "A continuacion se detalla como cada objetivo especifico fue abordado y cumplido, "
            "junto con la evidencia concreta que lo sustenta."
        ),
        first_line_indent=1.25,
    )

    add_styled_table(
        doc,
        headers=["Objetivo", "Como se Resolvio", "Evidencia / Entregable"],
        rows=[
            [
                "OE1: Fundamentacion teorica de CP en riesgo crediticio",
                "Revision de 19 fuentes (2005-2024). "
                "Identificacion de Split, Mondrian y CQR como variantes "
                "pertinentes. Analisis comparativo con metodos "
                "tradicionales (Platt, Isotonic, Bootstrap).",
                "Marco teorico (seccion 5) con 5 subsecciones. "
                "19 referencias IEEE. Identificacion del vacio: "
                "CP aplicada casi exclusivamente a PD, no a LGD/EAD.",
            ],
            [
                "OE2: Dataset analitico con split temporal OOT",
                "Construccion de 3 conjuntos OOT a partir de LendingClub. "
                "Eliminacion de 22 variables con data leakage. "
                "Feature engineering de 42 features.",
                "Train: 1.35M (2007-2017). Cal: 237K (2017). "
                "Test: 277K (2018-2020). Shift de default rate "
                "18.5% -> 22.0% confirma robustez temporal.",
            ],
            [
                "OE3: Modelo base PD calibrado",
                "CatBoost monotónico + Venn-Abers seleccionada "
                "automaticamente via politica multi-fold temporal "
                "(4 folds, criterio Brier + AUC-guard).",
                "AUC-ROC: 0.7127. ECE: 0.0067. Brier: 0.155. "
                "D2-Brier: +0.099 (positivo post-calibracion). "
                "KS: 0.313.",
            ],
            [
                "OE4: CP Mondrian sobre PD, LGD y EAD",
                "Mondrian por grado (A-G) para PD. "
                "4 variantes benchmark para LGD (adaptive gano). "
                "Split Conformal para EAD.",
                "PD: 92.42% cob., min grado 89.01%. "
                "LGD: 90.50% cob. (unica variante que pasa guardrails). "
                "EAD: 91.20% cob., R2=0.9999.",
            ],
            [
                "OE5: Impacto regulatorio IFRS 9",
                "Simulacion de ECL bajo 4 escenarios con intervalos CP. "
                "Clasificacion por etapas Stage 1-3. "
                "Propuesta de ancho CP como senal SICR.",
                "ECL: $1,001M (base) a $1,799M (severe), +79.7%. "
                "Rango de $798M cuantifica incertidumbre. "
                "Stage 2: 43.01% del portafolio.",
            ],
            [
                "OE6: Lineamientos de adopcion",
                "Backtesting mensual (35 meses). Gate de politica "
                "con 13 checks. Dashboard de monitoreo. "
                "Criterios de escalamiento y recalibracion.",
                "Framework de gobernanza con alertas por grado. "
                "Pipeline DVC de 26 etapas. "
                "Dashboard Streamlit con 27 paginas.",
            ],
        ],
        caption="Cumplimiento de objetivos especificos con evidencia.",
        table_num=12,
    )


# ── 10. CONCLUSIONES ─────────────────────────────────────────────────


def build_conclusions(doc: Document) -> None:
    styled_heading(doc, "10. CONCLUSIONES", level=1)

    add_paragraph(
        doc,
        (
            "Las conclusiones de esta investigacion se articulan como respuesta directa a la pregunta "
            "de investigacion y a cada uno de los seis objetivos especificos planteados en la seccion 3. "
            "La Tabla 12 (seccion 9.3) documenta la evidencia cuantitativa que sustenta cada conclusion; "
            "aqui se ofrece la sintesis interpretativa y el juicio sobre el alcance de los resultados."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "La fundamentacion teorica (OE1) permitio identificar que la prediccion conformal posee una "
            "propiedad unica entre los metodos de cuantificacion de incertidumbre: ofrece garantias "
            "finitas de cobertura sin asumir distribuciones parametricas, lo cual la hace especialmente "
            "adecuada para portafolios crediticios donde los supuestos de normalidad son frecuentemente "
            "violados. La revision de la literatura revelo, ademas, un vacio concreto: la aplicacion de "
            "CP se habia limitado casi exclusivamente a PD, sin extension a LGD ni EAD, lo que "
            "constituyo la oportunidad de contribucion de este trabajo."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "La construccion del dataset analitico (OE2) demostro que un diseno de validacion "
            "out-of-time riguroso es condicion necesaria para evaluar CP en riesgo crediticio. El shift "
            "de tasa de default entre entrenamiento (18.52%) y test (21.98%) confirma que las metricas "
            "reportadas reflejan rendimiento prospectivo genuino, no mera evaluacion in-sample. Sin "
            "esta disciplina temporal, las coberturas conformales serian artificialmente optimistas."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "El modelo base calibrado (OE3) confirmo una conclusion de alto valor practico: la "
            "calibracion probabilistica y la prediccion conformal son mecanismos complementarios, no "
            "sustitutos. La calibracion Venn-Abers corrigio el nivel probabilistico del modelo "
            "(D2-Brier de -0.222 a +0.099), pero por si sola no responde con que confianza se emite "
            "cada prediccion. La incorporacion de restricciones monotonicas en el modelo base asegura "
            "coherencia economica en las predicciones, lo que a su vez mejora la interpretabilidad "
            "de los intervalos conformales construidos sobre ellas."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "La aplicacion de Mondrian Conformal Prediction a la triada PD-LGD-EAD (OE4) constituye "
            "la contribucion central de esta investigacion. El enfoque Mondrian supero al Split "
            "Conformal global en +30 puntos porcentuales de cobertura minima por grado, eliminando "
            "la falsa sensacion de seguridad que proporciona una metrica promedio que oculta "
            "subcoberturas severas en segmentos criticos. La extension a LGD y EAD —con coberturas "
            "de 90.50% y 91.20% respectivamente— demuestra que el framework conformal es aplicable "
            "a los tres componentes de la perdida crediticia, cerrando el vacio identificado en la "
            "literatura."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "El analisis de impacto regulatorio (OE5) tradujo las coberturas estadisticas en "
            "informacion accionable para comites de riesgo. El rango de ECL entre $1,001M y $1,799M "
            "proporciona por primera vez una cuantificacion explicita de la incertidumbre en las "
            "provisiones IFRS 9 derivada directamente de los intervalos conformales, no de supuestos "
            "ad hoc. La propuesta del ancho del intervalo conformal como senal complementaria de SICR "
            "abre una via concreta para integrar CP en los procesos regulatorios existentes."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Los lineamientos de adopcion (OE6) demostraron la viabilidad operativa del framework "
            "a traves de un backtesting de 35 meses con alertas de cobertura por grado. El gate de "
            "politica (7/13 checks, 0 criticos) ilustra un enfoque de gobernanza que prioriza "
            "transparencia sobre aprobacion automatica, permitiendo al equipo de riesgo tomar "
            "decisiones informadas sobre cuando recalibrar."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Es necesario reconocer las limitaciones del trabajo. El supuesto de intercambiabilidad "
            "que sustenta la garantia de cobertura conformal puede debilitarse en periodos de crisis "
            "economica severa. Los intervalos para grados de alto riesgo (F, G) tienden a ser mas "
            "amplios, reduciendo su utilidad practica para decisiones granulares. El dataset, aunque "
            "extenso (1.86M prestamos), proviene de una unica plataforma de credito al consumo, lo "
            "que limita la generalizabilidad a otros productos o mercados."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "En respuesta a la pregunta de investigacion: las tecnicas de prediccion conformal "
            "mejoran efectivamente la cuantificacion de incertidumbre en modelos de riesgo crediticio. "
            "El valor fundamental no reside en mejorar la discriminacion o la calibracion del modelo "
            "—que son tareas resueltas por otros mecanismos— sino en proporcionar una capa adicional "
            "de informacion que convierte predicciones puntuales en predicciones con garantia de "
            "cobertura. Esta distincion es critica para la gestion del riesgo: un modelo puede estar "
            "bien calibrado en promedio y aun asi ser fragil para decisiones individuales. La "
            "prediccion conformal, y en particular la variante Mondrian, resuelve esta fragilidad con "
            "garantias formales que se mantienen por segmento de riesgo."
        ),
        first_line_indent=1.25,
    )


# ── 11. LINEAS FUTURAS ───────────────────────────────────────────────


def build_future_work(doc: Document) -> None:
    styled_heading(doc, "11. LINEAS FUTURAS DEL PROYECTO", level=1)

    add_paragraph(
        doc,
        (
            "A partir de los resultados obtenidos y las limitaciones identificadas, se proponen las "
            "siguientes lineas de continuacion, enmarcadas en el plan de estudios de maestria:"
        ),
        first_line_indent=1.25,
    )

    futures = [
        (
            "Prediccion conformal adaptativa (ACI)",
            (
                "Implementar metodos de Adaptive Conformal Inference para manejar la no "
                "estacionariedad de los datos financieros, ajustando los quantiles de manera "
                "dinamica ante cambios en la distribucion."
            ),
        ),
        (
            "Integracion con optimizacion de portafolio",
            (
                "Utilizar los intervalos conformales como conjuntos de incertidumbre (uncertainty "
                "sets) en formulaciones de optimizacion robusta, conectando la cuantificacion de "
                "incertidumbre con decisiones de asignacion de capital."
            ),
        ),
        (
            "Conformal prediction online",
            (
                "Desarrollar un pipeline de CP en streaming para decisiones de credito en tiempo "
                "real, integrando actualizaciones incrementales sin recalibracion completa."
            ),
        ),
        (
            "Extension del analisis IFRS 9",
            (
                "Evaluar el uso del ancho del intervalo conformal como senal primaria de SICR en "
                "portafolios bancarios reales, comparando su poder predictivo con senales "
                "tradicionales de deterioro."
            ),
        ),
        (
            "Validacion en otros mercados",
            (
                "Aplicar el framework a otros tipos de productos crediticios (hipotecas, tarjetas, "
                "prestamos corporativos) para evaluar la generalizabilidad de los hallazgos."
            ),
        ),
    ]

    for title, desc in futures:
        p = doc.add_paragraph()
        r = p.add_run(f"{title}: ")
        r.font.bold = True
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"
        r.font.color.rgb = DARK_GRAY
        r2 = p.add_run(desc)
        r2.font.size = Pt(12)
        r2.font.name = "Times New Roman"
        r2.font.color.rgb = DARK_GRAY
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)


# ── 12. REFERENCIAS ──────────────────────────────────────────────────


def build_references(doc: Document) -> None:
    styled_heading(doc, "12. REFERENCIAS", level=1)

    refs = [
        '[1] A. N. Angelopoulos y S. Bates, "A Gentle Introduction to Conformal Prediction '
        'and Distribution-Free Uncertainty Quantification," Foundations and Trends in Machine '
        "Learning, vol. 16, no. 4, pp. 494-591, 2023.",
        "[2] B. Baesens, D. Roesch y H. Scheule, Credit Risk Analytics: Measurement "
        "Techniques, Applications, and Examples in SAS. John Wiley & Sons, 2016.",
        '[3] R. F. Barber, E. J. Candes, A. Ramdas y R. J. Tibshirani, "Predictive '
        'Inference with the Jackknife+," The Annals of Statistics, vol. 49, no. 1, '
        "pp. 486-507, 2021.",
        '[4] Basel Committee on Banking Supervision, "International Convergence of Capital '
        'Measurement and Capital Standards: A Revised Framework," Bank for International '
        "Settlements (BIS), 2006.",
        '[5] T. Bellotti, "Reliable region predictions for automated credit scoring," en '
        "Proc. Workshop on Conformal and Probabilistic Prediction and Applications "
        "(COPA), PMLR, 2017.",
        '[6] Deloitte, "IFRS 9 and Expected Credit Loss: Modelling and Validation '
        'Challenges," Deloitte Technical Report, 2020.',
        '[7] European Banking Authority (EBA), "Guidelines on Loan Origination and '
        'Monitoring," EBA/GL/2020/06, 2020.',
        '[8] M. Fontana, G. Zeni y S. Vantini, "Conformal Prediction: A Unified Review of '
        'Theory and New Challenges," Bernoulli, vol. 29, no. 1, pp. 1-23, 2023.',
        '[9] International Accounting Standards Board (IASB), "IFRS 9 Financial '
        'Instruments," IFRS Foundation, 2014.',
        '[10] F. Javanmardi y V. Vovk, "Multip probability predictions for credit scoring '
        'with Venn-Abers predictors," en Conformal and Probabilistic Prediction and '
        "Applications (COPA), PMLR, 2023.",
        '[11] Lending Club, "Lending Club Loan Data 2007-2020 Q1," Kaggle Dataset, 2020. '
        "[Online]. Disponible: https://www.kaggle.com/datasets/ethon0426/lending-club-20072020q1",
        '[12] S. Lessmann, B. Baesens, H.-V. Seow y L. C. Thomas, "Benchmarking '
        'state-of-the-art classification algorithms for credit scoring," European Journal of '
        "Operational Research, vol. 247, no. 1, pp. 124-136, 2015.",
        '[13] H. Lu, N. Karimireddy, N. Ponomareva y A. Mirrokni, "Conformal Prediction '
        'in the Medical Domain: A Comprehensive Survey," IEEE Reviews in Biomedical '
        "Engineering, vol. 17, pp. 127-142, 2024.",
        '[14] MAPIE Contributors, "MAPIE: Model Agnostic Prediction Interval Estimator," '
        "2023. [Online]. Disponible: https://mapie.readthedocs.io/",
        '[15] A. Niculescu-Mizil y R. Caruana, "Predicting good probabilities with supervised '
        'learning," en Proc. 22nd ICML, pp. 625-632, 2005.',
        '[16] J. Platt, "Probabilistic Outputs for Support Vector Machines and Comparisons '
        'to Regularized Likelihood Methods," en Advances in Large Margin Classifiers. '
        "MIT Press, 1999.",
        '[17] Y. Romano, E. Patterson y E. Candes, "Conformalized Quantile Regression," en '
        "Advances in Neural Information Processing Systems (NeurIPS), vol. 32, 2019.",
        "[18] V. Vovk, A. Gammerman y G. Shafer, Algorithmic Learning in a Random World, "
        "2nd ed. Springer, 2022.",
        '[19] V. Vovk y I. Petej, "Venn-Abers Predictors," en Proc. 30th Conf. on '
        "Uncertainty in Artificial Intelligence (UAI), pp. 829-838, 2014.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(ref)
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        r.font.color.rgb = DARK_GRAY


# ── ANEXO: DECLARACION IA ────────────────────────────────────────────


def build_ai_declaration(doc: Document) -> None:
    doc.add_page_break()
    styled_heading(
        doc, "ANEXO: DECLARACION DE USO DE HERRAMIENTAS DE INTELIGENCIA ARTIFICIAL", level=1
    )

    add_paragraph(
        doc,
        (
            "Este documento tecnico ha sido preparado con la asistencia de herramientas de "
            "inteligencia artificial utilizadas estrictamente como ayudas de programacion y "
            "redaccion. Especificamente, Claude (Anthropic) fue empleado como asistente de "
            "programacion para la implementacion del pipeline de machine learning, la generacion "
            "de scripts de experimentacion y la estructuracion del documento tecnico final. Tambien "
            "fue utilizado como asistente de redaccion para mejorar la claridad, gramatica y estilo "
            "a lo largo del manuscrito."
        ),
        first_line_indent=1.25,
    )

    add_paragraph(
        doc,
        (
            "Todo el contenido intelectual y cientifico, incluyendo las ideas, analisis, resultados, "
            "interpretaciones y conclusiones presentadas en este trabajo, son responsabilidad "
            "exclusiva del autor. Las herramientas de inteligencia artificial sirvieron unicamente "
            "para facilitar la calidad linguistica y presentacional del documento, asi como para "
            "acelerar la implementacion tecnica del codigo fuente. Todas las decisiones "
            "metodologicas, la seleccion de tecnicas, el diseno experimental y las conclusiones "
            "fueron realizadas por el autor bajo su completa supervision y verificacion."
        ),
        first_line_indent=1.25,
    )


# ══════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════


def main() -> None:
    from loguru import logger

    logger.info("Generando Documento Tecnico Final (scope: Conformal Prediction)...")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = DARK_GRAY
    style.paragraph_format.line_spacing = 1.5

    fig_num = 1

    build_cover_page(doc)
    logger.info("  Portada")

    build_abstract(doc)
    logger.info("  Resumen / Abstract")

    build_table_of_contents(doc)
    logger.info("  Contenido")

    build_introduction(doc)
    doc.add_page_break()
    logger.info("  1. Introduccion")

    build_problem_statement(doc)
    doc.add_page_break()
    logger.info("  2. Planteamiento del Problema")

    build_justification(doc)
    doc.add_page_break()
    logger.info("  3. Justificacion")

    build_objectives(doc)
    doc.add_page_break()
    logger.info("  4. Objetivos")

    fig_num = build_theoretical_framework(doc, fig_num)
    doc.add_page_break()
    logger.info("  5. Marco Teorico")

    fig_num = build_methodology(doc, fig_num)
    doc.add_page_break()
    logger.info("  6. Metodologia")

    build_ethics(doc)
    doc.add_page_break()
    logger.info("  7. Consideraciones Eticas")

    fig_num = build_results(doc, fig_num)
    doc.add_page_break()
    logger.info("  8. Resultados")

    build_discussion(doc)
    doc.add_page_break()
    logger.info("  9. Discusion")

    build_conclusions(doc)
    doc.add_page_break()
    logger.info("  10. Conclusiones")

    build_future_work(doc)
    doc.add_page_break()
    logger.info("  11. Lineas Futuras")

    build_references(doc)
    logger.info("  12. Referencias")

    build_ai_declaration(doc)
    logger.info("  Anexo: Declaracion IA")

    doc.save(str(DOCX_PATH))
    logger.info(f"Documento guardado en: {DOCX_PATH}")
    logger.info(f"Total figuras: {fig_num - 1}")


if __name__ == "__main__":
    main()
